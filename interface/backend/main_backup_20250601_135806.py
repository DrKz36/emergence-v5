"""
ÉMERGENCE V5 - FastAPI Backend COMPLET - SUPPORT 3 MODES
🔥 Mode Dialogue + Mode Triangle + Mode Documents + Upload PDF/DOCX
FIX: WebSocket mapping + Mode Documents + Costs tracking + Responsive fixes
"""

import sys
import os
import hashlib
import tempfile
import mimetypes
from pathlib import Path
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File, HTTPException, Query
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Optional, Any
import json
import asyncio
import logging
from datetime import datetime
import uuid

# Imports ÉMERGENCE V4
try:
    from core.database import get_db
    from core.vector_manager import get_vector_manager
    from core.rag_manager import get_rag_manager
    from core.agents import get_real_agents, AgentResponse
    MODULES_AVAILABLE = True
    print("✅ Modules ÉMERGENCE V4 chargés")
except ImportError as e:
    print(f"❌ Erreur import modules V4: {e}")
    MODULES_AVAILABLE = False

# Imports extraction documents
try:
    import PyPDF2
    PDF_AVAILABLE = True
    print("✅ PyPDF2 disponible")
except ImportError:
    PDF_AVAILABLE = False
    print("❌ PyPDF2 non disponible")

try:
    from docx import Document
    DOCX_AVAILABLE = True
    print("✅ python-docx disponible")
except ImportError:
    DOCX_AVAILABLE = False
    print("❌ python-docx non disponible")

# Configuration logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# FastAPI app
app = FastAPI(
    title="ÉMERGENCE V5 - Multi-IA API 3 Modes",
    description="Interface V5: Mode Dialogue + Mode Triangle + Mode Documents avec sélection granulaire",
    version="5.0.0"
)

# CORS pour frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Servir fichiers statiques
app.mount("/static", StaticFiles(directory="interface/frontend"), name="static")

# === MODELS PYDANTIC V5 ===
class ChatMessage(BaseModel):
    agent: str
    message: str
    use_rag: bool = True
    rag_chunks: int = 5

class TripleMessage(BaseModel):
    message: str
    use_rag: bool = True
    rag_chunks: int = 5

class DocumentsMessage(BaseModel):
    """🆕 Nouveau modèle pour Mode Documents"""
    message: str
    agents: List[str] = ['anima']  # Liste des agents sélectionnés
    documents: List[str] = []      # Liste des IDs documents sélectionnés
    use_rag: bool = True
    rag_chunks: int = 5

class AgentResponseModel(BaseModel):
    agent: str
    response_text: str
    processing_time: float
    model_used: str
    rag_chunks_count: int
    cost_estimate: float
    provider: str
    timestamp: str

class SystemStatus(BaseModel):
    modules_available: bool
    agents_status: Dict[str, Any]
    database_stats: Dict[str, Any]
    total_cost: float

class DocumentInfo(BaseModel):
    """🗄️ Modèle information document"""
    id: str
    filename: str
    file_size: int
    upload_date: str
    chunks_count: int
    file_hash: str
    format: str
    text_preview: Optional[str] = ""
    processing_status: str = "completed"

class DocumentListResponse(BaseModel):
    """📋 Réponse liste documents"""
    documents: List[DocumentInfo]
    total: int
    total_size: int
    supported_formats: List[str]
    pdf_available: bool
    docx_available: bool

class DatabaseStats(BaseModel):
    """📊 Statistiques base de données"""
    documents_count: int
    chunks_count: int
    conversations_count: int
    total_size_mb: float
    database_path: str
    vector_count: int
    last_update: str

# === UTILITAIRES EXTRACTION DOCUMENTS (inchangés) ===
def extract_text_from_pdf(file_content: bytes) -> str:
    """🔴 Extrait le texte d'un PDF - Version optimisée"""
    if not PDF_AVAILABLE:
        raise ValueError("PyPDF2 non installé - impossible de lire les PDF")
    
    try:
        # Sauvegarde temporaire
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
            tmp_file.write(file_content)
            tmp_path = tmp_file.name
        
        # Extraction texte avec gestion d'erreurs robuste
        text_parts = []
        try:
            with open(tmp_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Vérification PDF valide
                if len(pdf_reader.pages) == 0:
                    raise ValueError("PDF vide ou corrompu")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(f"[Page {page_num + 1}]\n{page_text.strip()}")
                    except Exception as page_error:
                        logger.warning(f"⚠️ Erreur page {page_num + 1}: {page_error}")
                        continue
                        
        finally:
            # Nettoyage garanti
            try:
                os.unlink(tmp_path)
            except:
                pass
        
        if not text_parts:
            raise ValueError("Aucun texte extractible du PDF")
        
        return "\n\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"❌ Erreur extraction PDF: {e}")
        raise ValueError(f"Impossible d'extraire le texte du PDF: {e}")

def extract_text_from_docx(file_content: bytes) -> str:
    """📘 Extrait le texte d'un DOCX - Version optimisée"""
    if not DOCX_AVAILABLE:
        raise ValueError("python-docx non installé - impossible de lire les DOCX")
    
    try:
        # Sauvegarde temporaire
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            tmp_file.write(file_content)
            tmp_path = tmp_file.name
        
        # Extraction texte avec gestion structure
        text_parts = []
        try:
            doc = Document(tmp_path)
            
            # Extraction paragraphes
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extraction tables si présentes
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.append("[TABLE]\n" + "\n".join(table_text))
                    
        finally:
            # Nettoyage garanti
            try:
                os.unlink(tmp_path)
            except:
                pass
        
        if not text_parts:
            raise ValueError("Document DOCX vide ou sans texte extractible")
        
        return "\n\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"❌ Erreur extraction DOCX: {e}")
        raise ValueError(f"Impossible d'extraire le texte du DOCX: {e}")

def extract_text_from_file(filename: str, file_content: bytes) -> str:
    """📄 Extrait le texte selon le type de fichier - Version complète"""
    file_ext = Path(filename).suffix.lower()
    
    try:
        if file_ext == '.pdf':
            return extract_text_from_pdf(file_content)
        elif file_ext == '.docx':
            return extract_text_from_docx(file_content)
        elif file_ext in ['.txt', '.md']:
            # Détection encoding pour texte
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    return file_content.decode('latin-1')
                except UnicodeDecodeError:
                    return file_content.decode('utf-8', errors='ignore')
        elif file_ext == '.json':
            # Validation JSON + extraction texte
            try:
                json_data = json.loads(file_content.decode('utf-8'))
                return json.dumps(json_data, indent=2, ensure_ascii=False)
            except json.JSONDecodeError as e:
                raise ValueError(f"Fichier JSON invalide: {e}")
        else:
            raise ValueError(f"Format de fichier non supporté: {file_ext}")
            
    except ValueError:
        # Re-raise les erreurs de validation
        raise
    except Exception as e:
        raise ValueError(f"Erreur inattendue lors de l'extraction {file_ext}: {e}")

def get_supported_formats() -> List[str]:
    """📋 Retourne les formats supportés avec status"""
    formats = ['.txt', '.md', '.json']
    
    if PDF_AVAILABLE:
        formats.append('.pdf')
    if DOCX_AVAILABLE:
        formats.append('.docx')
    
    return formats

def validate_file_content(filename: str, file_content: bytes, extracted_text: str) -> Dict[str, Any]:
    """✅ Validation contenu fichier avec métriques"""
    file_ext = Path(filename).suffix.lower()
    
    validation = {
        'valid': True,
        'errors': [],
        'warnings': [],
        'metrics': {
            'file_size': len(file_content),
            'text_length': len(extracted_text),
            'lines_count': len(extracted_text.splitlines()),
            'words_count': len(extracted_text.split()),
            'estimated_chunks': 0
        }
    }
    
    # Validation taille minimum
    if len(extracted_text.strip()) < 10:
        validation['valid'] = False
        validation['errors'].append("Document contient moins de 10 caractères de texte")
    
    # Validation contenu selon format
    if file_ext == '.pdf' and validation['metrics']['text_length'] < 50:
        validation['warnings'].append("PDF semble avoir très peu de texte extractible")
    
    if file_ext == '.docx' and validation['metrics']['words_count'] < 5:
        validation['warnings'].append("Document DOCX semble très court")
    
    # Estimation chunks pour indexation
    paragraphs = [p.strip() for p in extracted_text.split('\n\n') if len(p.strip()) >= 50]
    validation['metrics']['estimated_chunks'] = len(paragraphs)
    
    if validation['metrics']['estimated_chunks'] == 0:
        validation['warnings'].append("Aucun chunk substantiel détecté pour indexation")
    
    return validation

# === GESTIONNAIRE CONNEXIONS WEBSOCKET V5 ===
class ConnectionManager:
    def __init__(self):
        self.active_connections: List[WebSocket] = []
        self.session_data: Dict[str, Any] = {}

    async def connect(self, websocket: WebSocket, session_id: str):
        await websocket.accept()
        self.active_connections.append(websocket)
        self.session_data[session_id] = {
            'websocket': websocket,
            'messages': [],
            'costs': {  # 🆕 Tracking coûts par agent
                'anima': 0.0,
                'neo': 0.0,
                'nexus': 0.0,
                'triple': 0.0,
                'total': 0.0
            }
        }
        logger.info(f"🔗 Nouvelle connexion WebSocket V5: {session_id}")

    def disconnect(self, websocket: WebSocket, session_id: str):
        if websocket in self.active_connections:
            self.active_connections.remove(websocket)
        if session_id in self.session_data:
            del self.session_data[session_id]
        logger.info(f"🔌 Connexion fermée: {session_id}")

    async def send_message(self, session_id: str, message: dict):
        if session_id in self.session_data:
            websocket = self.session_data[session_id]['websocket']
            try:
                await websocket.send_text(json.dumps(message))
            except Exception as e:
                logger.error(f"❌ Erreur envoi message: {e}")

    def update_cost(self, session_id: str, agent: str, cost: float):
        """🆕 Mise à jour coûts par agent"""
        if session_id in self.session_data:
            costs = self.session_data[session_id]['costs']
            costs[agent] += cost
            costs['total'] += cost

    def get_costs(self, session_id: str) -> Dict[str, float]:
        """🆕 Récupération coûts session"""
        if session_id in self.session_data:
            return self.session_data[session_id]['costs']
        return {'anima': 0.0, 'neo': 0.0, 'nexus': 0.0, 'triple': 0.0, 'total': 0.0}

manager = ConnectionManager()

# === ROUTES PRINCIPALES ===
@app.get("/", response_class=HTMLResponse)
async def read_root():
    """🏠 Sert la page principale ÉMERGENCE V5"""
    try:
        with open("interface/frontend/index.html", "r", encoding="utf-8") as f:
            html_content = f.read()
            
            # Injection info formats supportés
            formats_info = f"<!-- ÉMERGENCE V5 - Formats supportés: {', '.join(get_supported_formats())} -->"
            html_content = html_content.replace("</head>", f"{formats_info}\n</head>")
            
            return HTMLResponse(content=html_content)
    except FileNotFoundError:
        return HTMLResponse(f"""
        <html>
            <head><title>ÉMERGENCE V5 - Interface Required</title></head>
            <body>
                <h1>🚀 ÉMERGENCE V5 - Backend Running</h1>
                <p><strong>Formats supportés:</strong> {", ".join(get_supported_formats())}</p>
                <p><strong>Modules V4:</strong> {'✅ Disponibles' if MODULES_AVAILABLE else '❌ Non disponibles'}</p>
                <h2>🎯 Modes supportés:</h2>
                <ul>
                    <li><strong>Mode Dialogue</strong> - Chat avec agent unique</li>
                    <li><strong>Mode Triangle</strong> - Débat triangulaire 3 agents</li>
                    <li><strong>Mode Documents</strong> - Sélection granulaire docs + agents</li>
                </ul>
                <h2>API Endpoints:</h2>
                <ul>
                    <li><a href="/api/status">/api/status</a> - Status système</li>
                    <li><a href="/api/documents">/api/documents</a> - Liste documents</li>
                    <li><a href="/api/database/stats">/api/database/stats</a> - Stats BDD</li>
                </ul>
                <p>Ajoutez le fichier <code>interface/frontend/index.html</code> avec l'interface V5.</p>
            </body>
        </html>
        """)

@app.get("/api/status")
async def get_system_status() -> SystemStatus:
    """📊 Status du système V5"""
    try:
        if MODULES_AVAILABLE:
            real_agents = get_real_agents()
            agents_status = real_agents.get_health_status()
            
            rag_manager = get_rag_manager()
            db_stats = rag_manager.get_stats()
            
            # Calcul coût total robuste
            try:
                usage_stats = real_agents.get_usage_stats()
                total_cost = 0.0
                if isinstance(usage_stats, dict):
                    if 'totals' in usage_stats and isinstance(usage_stats['totals'], dict):
                        total_cost = usage_stats['totals'].get('cost', 0.0)
                    elif 'total_cost' in usage_stats:
                        total_cost = usage_stats['total_cost']
                    elif 'cost' in usage_stats:
                        total_cost = usage_stats['cost']
                    else:
                        for agent_stats in usage_stats.values():
                            if isinstance(agent_stats, dict) and 'cost' in agent_stats:
                                total_cost += agent_stats['cost']
            except Exception as e:
                logger.warning(f"⚠️ Erreur calcul coût total: {e}")
                total_cost = 0.0
            
            return SystemStatus(
                modules_available=True,
                agents_status=agents_status,
                database_stats=db_stats,
                total_cost=total_cost
            )
        else:
            return SystemStatus(
                modules_available=False,
                agents_status={},
                database_stats={},
                total_cost=0.0
            )
    except Exception as e:
        logger.error(f"❌ Erreur status: {e}")
        return SystemStatus(
            modules_available=False,
            agents_status={'error': str(e)},
            database_stats={'error': str(e)},
            total_cost=0.0
        )

# === ROUTES CHAT EXISTANTES (compatibilité) ===
@app.post("/api/chat")
async def chat_with_agent(message: ChatMessage) -> AgentResponseModel:
    """💬 Chat avec un agent individuel"""
    if not MODULES_AVAILABLE:
        raise HTTPException(status_code=503, detail="Modules V4 non disponibles")
    
    try:
        # Récupération contexte RAG
        context = ""
        rag_chunks_count = 0
        if message.use_rag:
            rag_manager = get_rag_manager()
            context = rag_manager.get_context_for_agent(
                message.message, 
                message.agent, 
                message.rag_chunks
            )
            rag_chunks_count = context.count("=== DOCUMENT") + context.count("--- Interaction")
        
        # Réponse agent
        real_agents = get_real_agents()
        response = real_agents.get_response(message.agent, message.message, context)
        
        # Stockage interaction V4
        try:
            rag_manager = get_rag_manager()
            session_id = f"api_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            rag_manager.store_interaction_v4(
                session_id=session_id,
                user_message=message.message,
                agent_name=message.agent,
                agent_response=response.response_text,
                processing_time=response.processing_time,
                rag_chunks_used=rag_chunks_count
            )
        except Exception as store_error:
            logger.warning(f"⚠️ Erreur stockage interaction (non-critique): {store_error}")
        
        return AgentResponseModel(
            agent=message.agent,
            response_text=response.response_text,
            processing_time=response.processing_time,
            model_used=response.model_used,
            rag_chunks_count=rag_chunks_count,
            cost_estimate=response.cost_estimate,
            provider=response.provider,
            timestamp=datetime.now().isoformat()
        )
        
    except Exception as e:
        logger.error(f"❌ Erreur chat agent {message.agent}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/triple")
async def triple_chat(message: TripleMessage) -> List[AgentResponseModel]:
    """🔺 Mode Triple - Les 3 agents débattent"""
    if not MODULES_AVAILABLE:
        raise HTTPException(status_code=503, detail="Modules V4 non disponibles")
    
    try:
        # Récupération contexte RAG
        context = ""
        rag_chunks_count = 0
        if message.use_rag:
            rag_manager = get_rag_manager()
            context = rag_manager.get_context_for_agent(
                message.message, 
                "Triple", 
                message.rag_chunks
            )
            rag_chunks_count = context.count("=== DOCUMENT") + context.count("--- Interaction")
        
        # Mode Triple
        real_agents = get_real_agents()
        triple_response = real_agents.get_triple_response(message.message, context)
        
        # Conversion en modèles API
        api_responses = []
        session_id = f"api_triple_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        if triple_response.success and triple_response.responses:
            for agent_name, response in triple_response.responses.items():
                if response and hasattr(response, 'response_text'):
                    api_responses.append(AgentResponseModel(
                        agent=agent_name.title(),
                        response_text=response.response_text,
                        processing_time=response.processing_time,
                        model_used=response.model_used,
                        rag_chunks_count=rag_chunks_count,
                        cost_estimate=response.cost_estimate,
                        provider=response.provider,
                        timestamp=datetime.now().isoformat()
                    ))
                    
                    # Stockage V4
                    try:
                        rag_manager = get_rag_manager()
                        rag_manager.store_interaction_v4(
                            session_id=session_id,
                            user_message=message.message,
                            agent_name=agent_name.title(),
                            agent_response=response.response_text,
                            processing_time=response.processing_time,
                            rag_chunks_used=rag_chunks_count
                        )
                    except Exception as store_error:
                        logger.warning(f"⚠️ Erreur stockage triple (non-critique): {store_error}")
        else:
            raise HTTPException(status_code=500, detail="Échec du mode Triple - aucune réponse générée")
        
        return api_responses
        
    except Exception as e:
        logger.error(f"❌ Erreur Triple mode: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# 🆕 NOUVELLE ROUTE - MODE DOCUMENTS
@app.post("/api/documents/chat")
async def documents_chat(message: DocumentsMessage) -> List[AgentResponseModel]:
    """📁 Mode Documents - Chat avec agents sélectionnés sur documents filtrés"""
    if not MODULES_AVAILABLE:
        raise HTTPException(status_code=503, detail="Modules V4 non disponibles")
    
    try:
        # Validation agents sélectionnés
        valid_agents = ['anima', 'neo', 'nexus']
        selected_agents = [agent for agent in message.agents if agent in valid_agents]
        
        if not selected_agents:
            raise HTTPException(status_code=400, detail="Aucun agent valide sélectionné")
        
        # Récupération contexte RAG filtré sur documents sélectionnés
        context = ""
        rag_chunks_count = 0
        if message.use_rag:
            rag_manager = get_rag_manager()
            
            if message.documents:
                # 🔥 NOUVEAU - RAG filtré sur documents spécifiques
                try:
                    if hasattr(rag_manager, 'get_context_filtered_documents'):
                        context = rag_manager.get_context_filtered_documents(
                            message.message, 
                            message.documents, 
                            message.rag_chunks
                        )
                    else:
                        # Fallback vers méthode standard
                        context = rag_manager.get_context_for_agent(
                            message.message, 
                            "Documents", 
                            message.rag_chunks
                        )
                except Exception as rag_error:
                    logger.warning(f"⚠️ Erreur RAG filtré, fallback standard: {rag_error}")
                    context = rag_manager.get_context_for_agent(
                        message.message, 
                        "Documents", 
                        message.rag_chunks
                    )
            else:
                # RAG sur tous documents si aucun sélectionné
                context = rag_manager.get_context_for_agent(
                    message.message, 
                    "Documents", 
                    message.rag_chunks
                )
            
            rag_chunks_count = context.count("=== DOCUMENT") + context.count("--- Interaction")
        
        # Réponses multiples agents
        api_responses = []
        session_id = f"api_docs_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        real_agents = get_real_agents()
        
        for agent_name in selected_agents:
            try:
                response = real_agents.get_response(agent_name, message.message, context)
                
                if response and hasattr(response, 'response_text'):
                    api_responses.append(AgentResponseModel(
                        agent=agent_name,
                        response_text=response.response_text,
                        processing_time=response.processing_time,
                        model_used=response.model_used,
                        rag_chunks_count=rag_chunks_count,
                        cost_estimate=response.cost_estimate,
                        provider=response.provider,
                        timestamp=datetime.now().isoformat()
                    ))
                    
                    # Stockage V4
                    try:
                        rag_manager = get_rag_manager()
                        rag_manager.store_interaction_v4(
                            session_id=session_id,
                            user_message=message.message,
                            agent_name=agent_name,
                            agent_response=response.response_text,
                            processing_time=response.processing_time,
                            rag_chunks_used=rag_chunks_count
                        )
                    except Exception as store_error:
                        logger.warning(f"⚠️ Erreur stockage documents mode (non-critique): {store_error}")
                        
            except Exception as agent_error:
                logger.error(f"❌ Erreur agent {agent_name} en mode documents: {agent_error}")
                # Continue avec les autres agents même si un échoue
                continue
        
        if not api_responses:
            raise HTTPException(status_code=500, detail="Aucune réponse générée par les agents sélectionnés")
        
        return api_responses
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Erreur mode Documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# === UPLOAD + GESTION DOCUMENTS (inchangé) ===
@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...)) -> dict:
    """🔥 Upload et indexation de document - SUPPORT COMPLET PDF/DOCX/TXT/MD/JSON"""
    if not MODULES_AVAILABLE:
        raise HTTPException(status_code=503, detail="Modules V4 non disponibles")
    
    try:
        # Validation taille fichier (max 10MB)
        MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
        file_content = await file.read()
        
        if len(file_content) > MAX_FILE_SIZE:
            raise HTTPException(status_code=413, detail=f"Fichier trop volumineux (max 10MB)")
        
        if len(file_content) == 0:
            raise HTTPException(status_code=400, detail="Fichier vide")
        
        # Validation extension
        file_ext = Path(file.filename).suffix.lower()
        supported_formats = get_supported_formats()
        
        if file_ext not in supported_formats:
            raise HTTPException(
                status_code=400, 
                detail=f"Format non supporté: {file_ext}. Formats acceptés: {', '.join(supported_formats)}"
            )
        
        # Extraction texte selon format
        try:
            text_content = extract_text_from_file(file.filename, file_content)
        except ValueError as ve:
            raise HTTPException(status_code=400, detail=str(ve))
        
        # Validation contenu avec métriques
        validation = validate_file_content(file.filename, file_content, text_content)
        
        if not validation['valid']:
            raise HTTPException(status_code=400, detail=f"Validation échouée: {'; '.join(validation['errors'])}")
        
        # Hash fichier pour éviter doublons
        file_hash = hashlib.md5(file_content).hexdigest()
        
        # Indexation avec RAG Manager V4
        rag_manager = get_rag_manager()
        success = rag_manager.store_document_v4(
            filename=file.filename,
            content=text_content,
            file_path=f"uploads/{file.filename}",
            file_hash=file_hash,
            file_size=len(file_content)
        )
        
        if success:
            response_data = {
                'success': True,
                'filename': file.filename,
                'format': file_ext,
                'file_hash': file_hash,
                'validation': validation,
                'message': f"Document '{file.filename}' indexé avec succès"
            }
            
            # Ajout warnings si présents
            if validation['warnings']:
                response_data['warnings'] = validation['warnings']
            
            return response_data
        else:
            raise HTTPException(status_code=500, detail="Erreur lors de l'indexation en base")
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Erreur upload: {e}")
        raise HTTPException(status_code=500, detail=f"Erreur inattendue: {str(e)}")

@app.get("/api/documents")
async def list_documents(
    limit: int = Query(50, ge=1, le=100, description="Nombre max de documents à retourner"),
    offset: int = Query(0, ge=0, description="Offset pour pagination")
) -> DocumentListResponse:
    """📋 Liste des documents indexés avec détails complets"""
    if not MODULES_AVAILABLE:
        return DocumentListResponse(
            documents=[],
            total=0,
            total_size=0,
            supported_formats=get_supported_formats(),
            pdf_available=PDF_AVAILABLE,
            docx_available=DOCX_AVAILABLE
        )
    
    try:
        # Implémentation via database directe
        db = get_db()
        documents = []
        total_docs = 0
        total_size = 0
        
        # Récupération documents via SQL direct
        try:
            # Query pour documents avec pagination
            with db.get_connection() as conn:
                cursor = conn.cursor()
                
                # Count total
                cursor.execute("SELECT COUNT(*) FROM documents")
                total_docs = cursor.fetchone()[0]
                
                # Documents avec détails
                cursor.execute("""
                    SELECT 
                        id, filename, file_size, upload_date, 
                        chunks_count, file_hash, processing_status, metadata
                    FROM documents 
                    ORDER BY upload_date DESC 
                    LIMIT ? OFFSET ?
                """, (limit, offset))
                
                for row in cursor.fetchall():
                    doc_id, filename, file_size, upload_date, chunks_count, file_hash, status, metadata = row
                    
                    # Format détection
                    file_format = Path(filename).suffix.lower() if filename else 'unknown'
                    
                    # Text preview via chunks
                    text_preview = ""
                    try:
                        cursor.execute("""
                            SELECT content FROM chunks 
                            WHERE document_id = ? 
                            ORDER BY chunk_index LIMIT 1
                        """, (doc_id,))
                        chunk_row = cursor.fetchone()
                        if chunk_row and chunk_row[0]:
                            preview_text = chunk_row[0][:200]
                            text_preview = preview_text + "..." if len(chunk_row[0]) > 200 else preview_text
                    except:
                        pass
                    
                    documents.append(DocumentInfo(
                        id=doc_id,
                        filename=filename or 'unknown',
                        file_size=file_size or 0,
                        upload_date=upload_date or '',
                        chunks_count=chunks_count or 0,
                        file_hash=file_hash or '',
                        format=file_format,
                        text_preview=text_preview,
                        processing_status=status or 'completed'
                    ))
                    
                    total_size += (file_size or 0)
                
        except Exception as db_error:
            logger.error(f"❌ Erreur requête documents database: {db_error}")
            # Fallback via RAG stats si database fail
            rag_manager = get_rag_manager()
            stats = rag_manager.get_stats()
            total_docs = stats.get('database', {}).get('documents_count', 0)
        
        return DocumentListResponse(
            documents=documents,
            total=total_docs,
            total_size=total_size,
            supported_formats=get_supported_formats(),
            pdf_available=PDF_AVAILABLE,
            docx_available=DOCX_AVAILABLE
        )
        
    except Exception as e:
        logger.error(f"❌ Erreur liste documents: {e}")
        # Retour minimal en cas d'erreur
        return DocumentListResponse(
            documents=[],
            total=0,
            total_size=0,
            supported_formats=get_supported_formats(),
            pdf_available=PDF_AVAILABLE,
            docx_available=DOCX_AVAILABLE
        )

@app.delete("/api/documents/{document_id}")
async def delete_document(document_id: str) -> dict:
    """🗑️ Suppression sélective d'un document"""
    if not MODULES_AVAILABLE:
        raise HTTPException(status_code=503, detail="Modules V4 non disponibles")
    
    try:
        # Implémentation via database directe
        db = get_db()
        
        # Vérification existence + récupération info
        with db.get_connection() as conn:
            cursor = conn.cursor()
            
            # Check existence
            cursor.execute("SELECT filename FROM documents WHERE id = ?", (document_id,))
            doc_row = cursor.fetchone()
            
            if not doc_row:
                raise HTTPException(status_code=404, detail="Document non trouvé")
            
            filename = doc_row[0]
            
            # Suppression chunks associés (cascade normalement gérée par FK)
            cursor.execute("DELETE FROM chunks WHERE document_id = ?", (document_id,))
            chunks_deleted = cursor.rowcount
            
            # Suppression document
            cursor.execute("DELETE FROM documents WHERE id = ?", (document_id,))
            doc_deleted = cursor.rowcount
            
            # Commit
            conn.commit()
            
            if doc_deleted > 0:
                logger.info(f"✅ Document supprimé: {filename} (ID: {document_id}, {chunks_deleted} chunks)")
                
                # TODO: Suppression vecteurs ChromaDB si possible
                try:
                    vector_manager = get_vector_manager()
                    if hasattr(vector_manager, 'delete_document_vectors'):
                        vector_manager.delete_document_vectors(document_id)
                except Exception as vector_error:
                    logger.warning(f"⚠️ Impossible de supprimer vecteurs: {vector_error}")
                
                return {
                    'success': True,
                    'document_id': document_id,
                    'filename': filename,
                    'chunks_deleted': chunks_deleted,
                    'message': f"Document '{filename}' supprimé avec succès"
                }
            else:
                raise HTTPException(status_code=500, detail="Échec suppression database")
                
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Erreur suppression document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Erreur inattendue: {str(e)}")

@app.get("/api/database/stats")
async def get_database_stats() -> DatabaseStats:
    """📊 Statistiques détaillées base de données"""
    if not MODULES_AVAILABLE:
        raise HTTPException(status_code=503, detail="Modules V4 non disponibles")
    
    try:
        # Stats via RAG Manager
        rag_manager = get_rag_manager()
        stats = rag_manager.get_stats()
        
        # Extraction métriques
        db_stats = stats.get('database', {})
        vector_stats = stats.get('vectors', {})
        
        # Calcul taille base (estimation)
        db_path = "data/emergence_v4.db"
        total_size_mb = 0.0
        try:
            if os.path.exists(db_path):
                total_size_mb = os.path.getsize(db_path) / (1024 * 1024)
        except:
            pass
        
        return DatabaseStats(
            documents_count=db_stats.get('documents_count', 0),
            chunks_count=db_stats.get('chunks_count', 0),
            conversations_count=db_stats.get('conversations_count', 0),
            total_size_mb=round(total_size_mb, 2),
            database_path=db_path,
            vector_count=vector_stats.get('vector_count', 0),
            last_update=datetime.now().isoformat()
        )
        
    except Exception as e:
        logger.error(f"❌ Erreur stats database: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# === WEBSOCKET V5 - SUPPORT 3 MODES ===
@app.websocket("/ws/{session_id}")
async def websocket_endpoint(websocket: WebSocket, session_id: str):
    """🔌 WebSocket ÉMERGENCE V5 - Support 3 modes + coûts tracking"""
    await manager.connect(websocket, session_id)
    
    try:
        while True:
            data = await websocket.receive_text()
            message_data = json.loads(data)
            
            msg_type = message_data.get('type', 'chat')
            
            if msg_type == 'chat':
                await handle_websocket_chat(session_id, message_data)
            elif msg_type == 'triple':
                await handle_websocket_triple(session_id, message_data)
            elif msg_type == 'documents':  # 🆕 Nouveau handler
                await handle_websocket_documents_mode(session_id, message_data)
            elif msg_type == 'status':
                await handle_websocket_status(session_id)
            elif msg_type == 'documents_list':
                await handle_websocket_documents_list(session_id, message_data)
                
    except WebSocketDisconnect:
        manager.disconnect(websocket, session_id)
    except Exception as e:
        logger.error(f"❌ Erreur WebSocket: {e}")
        try:
            await manager.send_message(session_id, {
                'type': 'error',
                'message': str(e)
            })
        except:
            pass

# === WEBSOCKET HANDLERS V5 ===
async def handle_websocket_chat(session_id: str, message_data: dict):
    """💬 Gestion chat WebSocket individuel"""
    try:
        if not MODULES_AVAILABLE:
            await manager.send_message(session_id, {
                'type': 'error',
                'message': 'Modules V4 non disponibles'
            })
            return
        
        agent = message_data.get('agent', 'anima')
        user_message = message_data.get('message', '')
        rag_enabled = message_data.get('rag_enabled', True)
        chunks_limit = message_data.get('chunks_limit', 5)
        
        await manager.send_message(session_id, {
            'type': 'typing',
            'agent': agent
        })
        
        # Contexte RAG
        context = ""
        rag_chunks_count = 0
        if rag_enabled:
            rag_manager = get_rag_manager()
            context = rag_manager.get_context_for_agent(user_message, agent, chunks_limit)
            rag_chunks_count = context.count("=== DOCUMENT") + context.count("--- Interaction")
        
        # Réponse agent
        real_agents = get_real_agents()
        response = real_agents.get_response(agent, user_message, context)
        
        # Stockage V4
        try:
            rag_manager = get_rag_manager()
            rag_manager.store_interaction_v4(
                session_id=f"ws_{session_id}",
                user_message=user_message,
                agent_name=agent,
                agent_response=response.response_text,
                processing_time=response.processing_time,
                rag_chunks_used=rag_chunks_count
            )
        except Exception as store_error:
            logger.warning(f"⚠️ Erreur stockage WebSocket (non-critique): {store_error}")
        
        # 🆕 Mise à jour coûts par agent
        manager.update_cost(session_id, agent, response.cost_estimate)
        session_costs = manager.get_costs(session_id)
        
        # Réponse WebSocket avec coûts
        await manager.send_message(session_id, {
            'type': 'agent_response',
            'agent': agent,
            'response': response.response_text,
            'metadata': {
                'model': response.model_used,
                'processing_time': response.processing_time,
                'rag_chunks': rag_chunks_count,
                'cost_estimate': response.cost_estimate,
                'provider': response.provider
            },
            'costs': session_costs,  # 🆕 Coûts session
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logger.error(f"❌ Erreur WebSocket chat: {e}")
        await manager.send_message(session_id, {
            'type': 'error',
            'message': str(e)
        })

async def handle_websocket_triple(session_id: str, message_data: dict):
    """🔺 Gestion mode Triple WebSocket"""
    try:
        if not MODULES_AVAILABLE:
            await manager.send_message(session_id, {
                'type': 'error',
                'message': 'Modules V4 non disponibles'
            })
            return
        
        user_message = message_data.get('message', '')
        rag_enabled = message_data.get('rag_enabled', True)
        chunks_limit = message_data.get('chunks_limit', 5)
        
        await manager.send_message(session_id, {
            'type': 'triple_start',
            'message': 'Débat triangulaire en cours...'
        })
        
        # Contexte RAG
        context = ""
        rag_chunks_count = 0
        if rag_enabled:
            rag_manager = get_rag_manager()
            context = rag_manager.get_context_for_agent(user_message, "Triple", chunks_limit)
            rag_chunks_count = context.count("=== DOCUMENT") + context.count("--- Interaction")
        
        # Mode Triple
        real_agents = get_real_agents()
        triple_response = real_agents.get_triple_response(user_message, context)
        
        # Envoi séquentiel des réponses
        total_cost = 0.0
        ws_session_id = f"ws_triple_{session_id}_{datetime.now().strftime('%H%M%S')}"
        
        if triple_response.success and triple_response.responses:
            for agent_name, response in triple_response.responses.items():
                if response and hasattr(response, 'response_text'):
                    total_cost += response.cost_estimate
                    
                    # Mise à jour coût session
                    manager.update_cost(session_id, 'triple', response.cost_estimate)
                    
                    await manager.send_message(session_id, {
                        'type': 'agent_response',
                        'agent': agent_name.lower(),
                        'response': response.response_text,
                        'metadata': {
                            'model': response.model_used,
                            'processing_time': response.processing_time,
                            'rag_chunks': rag_chunks_count,
                            'cost_estimate': response.cost_estimate,
                            'provider': response.provider
                        },
                        'timestamp': datetime.now().isoformat(),
                        'is_triple': True
                    })
                    
                    # Stockage V4
                    try:
                        rag_manager = get_rag_manager()
                        rag_manager.store_interaction_v4(
                            session_id=ws_session_id,
                            user_message=user_message,
                            agent_name=agent_name.title(),
                            agent_response=response.response_text,
                            processing_time=response.processing_time,
                            rag_chunks_used=rag_chunks_count
                        )
                    except Exception as store_error:
                        logger.warning(f"⚠️ Erreur stockage Triple WebSocket (non-critique): {store_error}")
        else:
            await manager.send_message(session_id, {
                'type': 'error',
                'message': 'Échec du mode Triple - aucune réponse générée'
            })
            return
        
        # Fin débat avec coûts
        session_costs = manager.get_costs(session_id)
        await manager.send_message(session_id, {
            'type': 'triple_end',
            'total_cost': total_cost,
            'agents_count': len(triple_response.responses) if triple_response.responses else 0,
            'costs': session_costs
        })
        
    except Exception as e:
        logger.error(f"❌ Erreur WebSocket Triple: {e}")
        await manager.send_message(session_id, {
            'type': 'error',
            'message': str(e)
        })

# 🆕 NOUVEAU HANDLER - MODE DOCUMENTS WEBSOCKET
async def handle_websocket_documents_mode(session_id: str, message_data: dict):
    """📁 Gestion mode Documents WebSocket - NOUVEAU"""
    try:
        if not MODULES_AVAILABLE:
            await manager.send_message(session_id, {
                'type': 'error',
                'message': 'Modules V4 non disponibles'
            })
            return
        
        user_message = message_data.get('message', '')
        agents = message_data.get('agents', ['anima'])
        documents = message_data.get('documents', [])
        rag_enabled = message_data.get('rag_enabled', True)
        chunks_limit = message_data.get('chunks_limit', 5)
        
        # Validation agents
        valid_agents = ['anima', 'neo', 'nexus']
        selected_agents = [agent for agent in agents if agent in valid_agents]
        
        if not selected_agents:
            await manager.send_message(session_id, {
                'type': 'error',
                'message': 'Aucun agent valide sélectionné'
            })
            return
        
        await manager.send_message(session_id, {
            'type': 'documents_start',
            'agents': selected_agents,
            'documents': documents,
            'message': f'Mode Documents: {len(selected_agents)} agents sur {len(documents)} documents'
        })
        
        # Contexte RAG filtré
        context = ""
        rag_chunks_count = 0
        if rag_enabled:
            rag_manager = get_rag_manager()
            
            if documents:
                # RAG filtré sur documents spécifiques
                try:
                    if hasattr(rag_manager, 'get_context_filtered_documents'):
                        context = rag_manager.get_context_filtered_documents(
                            user_message, 
                            documents, 
                            chunks_limit
                        )
                    else:
                        # Fallback vers méthode standard
                        context = rag_manager.get_context_for_agent(
                            user_message, 
                            "Documents", 
                            chunks_limit
                        )
                except Exception as rag_error:
                    logger.warning(f"⚠️ Erreur RAG filtré, fallback standard: {rag_error}")
                    context = rag_manager.get_context_for_agent(
                        user_message, 
                        "Documents", 
                        chunks_limit
                    )
            else:
                # RAG sur tous documents si aucun sélectionné
                context = rag_manager.get_context_for_agent(
                    user_message, 
                    "Documents", 
                    chunks_limit
                )
            
            rag_chunks_count = context.count("=== DOCUMENT") + context.count("--- Interaction")
        
        # Réponses multiples agents séquentielles
        real_agents = get_real_agents()
        ws_session_id = f"ws_docs_{session_id}_{datetime.now().strftime('%H%M%S')}"
        total_cost = 0.0
        
        for agent_name in selected_agents:
            try:
                await manager.send_message(session_id, {
                    'type': 'typing',
                    'agent': agent_name
                })
                
                response = real_agents.get_response(agent_name, user_message, context)
                
                if response and hasattr(response, 'response_text'):
                    total_cost += response.cost_estimate
                    
                    # Mise à jour coût session
                    manager.update_cost(session_id, agent_name, response.cost_estimate)
                    
                    await manager.send_message(session_id, {
                        'type': 'agent_response',
                        'agent': agent_name,
                        'response': response.response_text,
                        'metadata': {
                            'model': response.model_used,
                            'processing_time': response.processing_time,
                            'rag_chunks': rag_chunks_count,
                            'cost_estimate': response.cost_estimate,
                            'provider': response.provider
                        },
                        'timestamp': datetime.now().isoformat(),
                        'is_documents_mode': True,
                        'documents_context': {
                            'selected_documents': documents,
                            'agents_count': len(selected_agents)
                        }
                    })
                    
                    # Stockage V4
                    try:
                        rag_manager = get_rag_manager()
                        rag_manager.store_interaction_v4(
                            session_id=ws_session_id,
                            user_message=user_message,
                            agent_name=agent_name,
                            agent_response=response.response_text,
                            processing_time=response.processing_time,
                            rag_chunks_used=rag_chunks_count
                        )
                    except Exception as store_error:
                        logger.warning(f"⚠️ Erreur stockage documents mode (non-critique): {store_error}")
                        
            except Exception as agent_error:
                logger.error(f"❌ Erreur agent {agent_name} en mode documents: {agent_error}")
                await manager.send_message(session_id, {
                    'type': 'agent_error',
                    'agent': agent_name,
                    'error': str(agent_error)
                })
                # Continue avec les autres agents
                continue
        
        # Fin mode documents avec résumé
        session_costs = manager.get_costs(session_id)
        await manager.send_message(session_id, {
            'type': 'documents_end',
            'total_cost': total_cost,
            'agents_processed': len(selected_agents),
            'documents_used': len(documents),
            'costs': session_costs
        })
        
    except Exception as e:
        logger.error(f"❌ Erreur WebSocket mode Documents: {e}")
        await manager.send_message(session_id, {
            'type': 'error',
            'message': str(e)
        })

async def handle_websocket_status(session_id: str):
    """📊 Envoi status système WebSocket avec coûts session"""
    try:
        status = await get_system_status()
        session_costs = manager.get_costs(session_id)
        
        await manager.send_message(session_id, {
            'type': 'status',
            'system_status': status.dict(),
            'session_costs': session_costs
        })
        
    except Exception as e:
        logger.error(f"❌ Erreur WebSocket status: {e}")
        await manager.send_message(session_id, {
            'type': 'error',
            'message': f"Erreur status: {e}"
        })

async def handle_websocket_documents_list(session_id: str, message_data: dict):
    """🗄️ Gestion documents WebSocket"""
    try:
        action = message_data.get('action', 'list')
        
        if action == 'list':
            docs_response = await list_documents()
            await manager.send_message(session_id, {
                'type': 'documents_list',
                'documents': [doc.dict() for doc in docs_response.documents],
                'total': docs_response.total,
                'supported_formats': docs_response.supported_formats
            })
            
        elif action == 'delete':
            document_id = message_data.get('document_id')
            if document_id:
                try:
                    result = await delete_document(document_id)
                    await manager.send_message(session_id, {
                        'type': 'document_deleted',
                        'success': True,
                        'document_id': document_id,
                        'message': result.get('message', 'Document supprimé')
                    })
                except HTTPException as e:
                    await manager.send_message(session_id, {
                        'type': 'document_deleted',
                        'success': False,
                        'document_id': document_id,
                        'error': e.detail
                    })
            else:
                await manager.send_message(session_id, {
                    'type': 'error',
                    'message': 'document_id requis pour suppression'
                })
                
    except Exception as e:
        logger.error(f"❌ Erreur WebSocket documents: {e}")
        await manager.send_message(session_id, {
            'type': 'error',
            'message': str(e)
        })

# === POINT D'ENTRÉE ===
if __name__ == "__main__":
    import uvicorn
    
    print("🚀 ÉMERGENCE V5 - BACKEND COMPLET")
    print("="*50)
    print(f"📄 Formats supportés: {', '.join(get_supported_formats())}")
    print(f"🔴 PDF: {'✅' if PDF_AVAILABLE else '❌'}")
    print(f"📘 DOCX: {'✅' if DOCX_AVAILABLE else '❌'}")
    print(f"🔧 Modules V4: {'✅' if MODULES_AVAILABLE else '❌'}")
    print()
    print("🎯 MODES SUPPORTÉS:")
    print("   💬 Mode Dialogue - Chat agent unique")
    print("   🔀 Mode Triangle - Débat triangulaire 3 agents")
    print("   📁 Mode Documents - Agents multiples + docs filtrés")
    print()
    print("🆕 NOUVELLES FONCTIONNALITÉS V5:")
    print("   - Tracking coûts par agent en temps réel")
    print("   - Sélection granulaire documents + agents")
    print("   - WebSocket handler mode Documents")
    print("   - API /api/documents/chat pour mode Documents")
    print()
    print("🌐 Démarrage serveur...")
    
    uvicorn.run(
        app, 
        host="127.0.0.1", 
        port=8000,
        reload=True,
        log_level="info"
    )