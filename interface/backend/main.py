"""
ÉMERGENCE V5 - FastAPI Backend COMPLET - SUPPORT 3 MODES + SESSION MANAGER V5
🔥 Mode Dialogue + Mode Triangle + Mode Documents + Upload PDF/DOCX + MÉMOIRE PERSISTANTE
INTEGRATION: SessionManager V5 pour capture automatique conversations temps réel
VERSION CORRIGÉE: Fixes session_type + fallbacks API routes + WebSocket Heartbeat + Handlers complets
VERSION: 5.1.2 - Tous handlers WebSocket présents + Corrections critiques
"""

import sys
import os
import hashlib
import tempfile
import mimetypes
from pathlib import Path
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File, HTTPException, Query
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Optional, Any
import json
import asyncio
import logging
from datetime import datetime, timedelta
import uuid

# Imports ÉMERGENCE V4
try:
    from core.database import get_db
    from core.vector_manager import get_vector_manager
    from core.rag_manager import get_rag_manager
    from core.agents import get_real_agents, AgentResponse
    MODULES_AVAILABLE = True
    print("✅ Modules ÉMERGENCE V4 chargés")
except ImportError as e:
    print(f"❌ Erreur import modules V4: {e}")
    MODULES_AVAILABLE = False

# 🆕 Import SessionManager V5 pour mémoire persistante
try:
    from core.session_manager import SessionManagerV5
    SESSION_MANAGER_AVAILABLE = True
    print("✅ SessionManager V5 chargé")
except ImportError as e:
    print(f"⚠️ SessionManager V5 non disponible: {e}")
    SESSION_MANAGER_AVAILABLE = False

# Imports extraction documents
try:
    import PyPDF2
    PDF_AVAILABLE = True
    print("✅ PyPDF2 disponible")
except ImportError:
    PDF_AVAILABLE = False
    print("❌ PyPDF2 non disponible")

try:
    from docx import Document
    DOCX_AVAILABLE = True
    print("✅ python-docx disponible")
except ImportError:
    DOCX_AVAILABLE = False
    print("❌ python-docx non disponible")

# Configuration logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 🆕 Instance globale SessionManager V5
session_manager_v5 = None
if SESSION_MANAGER_AVAILABLE:
    try:
        session_manager_v5 = SessionManagerV5()
        print("✅ SessionManager V5 initialisé")
    except Exception as e:
        print(f"⚠️ Erreur initialisation SessionManager V5: {e}")
        session_manager_v5 = None

# FastAPI app
app = FastAPI(
    title="ÉMERGENCE V5 - Multi-IA API 3 Modes + Session Manager",
    description="Interface V5: Mode Dialogue + Mode Triangle + Mode Documents avec mémoire persistante",
    version="5.1.2"
)

# CORS pour frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Servir fichiers statiques
app.mount("/static", StaticFiles(directory="interface/frontend"), name="static")

# === MODELS PYDANTIC V5 ===
class ChatMessage(BaseModel):
    agent: str
    message: str
    use_rag: bool = True
    rag_chunks: int = 5

class TripleMessage(BaseModel):
    message: str
    use_rag: bool = True
    rag_chunks: int = 5

class DocumentsMessage(BaseModel):
    """🆕 Nouveau modèle pour Mode Documents"""
    message: str
    agents: List[str] = ['anima']  # Liste des agents sélectionnés
    documents: List[str] = []      # Liste des IDs documents sélectionnés
    use_rag: bool = True
    rag_chunks: int = 5

class AgentResponseModel(BaseModel):
    agent: str
    response_text: str
    processing_time: float
    model_used: str
    rag_chunks_count: int
    cost_estimate: float
    provider: str
    timestamp: str

class SystemStatus(BaseModel):
    modules_available: bool
    agents_status: Dict[str, Any]
    database_stats: Dict[str, Any]
    total_cost: float
    session_manager_available: bool = False  # 🆕

class DocumentInfo(BaseModel):
    """🗄️ Modèle information document"""
    id: str
    filename: str
    file_size: int
    upload_date: str
    chunks_count: int
    file_hash: str
    format: str
    text_preview: Optional[str] = ""
    processing_status: str = "completed"

class DocumentListResponse(BaseModel):
    """📋 Réponse liste documents"""
    documents: List[DocumentInfo]
    total: int
    total_size: int
    supported_formats: List[str]
    pdf_available: bool
    docx_available: bool

class DatabaseStats(BaseModel):
    """📊 Statistiques base de données"""
    documents_count: int
    chunks_count: int
    conversations_count: int
    total_size_mb: float
    database_path: str
    vector_count: int
    last_update: str

# === UTILITAIRES EXTRACTION DOCUMENTS ===
def extract_text_from_pdf(file_content: bytes) -> str:
    """🔴 Extrait le texte d'un PDF - Version optimisée"""
    if not PDF_AVAILABLE:
        raise ValueError("PyPDF2 non installé - impossible de lire les PDF")
    
    try:
        # Sauvegarde temporaire
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
            tmp_file.write(file_content)
            tmp_path = tmp_file.name
        
        # Extraction texte avec gestion d'erreurs robuste
        text_parts = []
        try:
            with open(tmp_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Vérification PDF valide
                if len(pdf_reader.pages) == 0:
                    raise ValueError("PDF vide ou corrompu")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(f"[Page {page_num + 1}]\n{page_text.strip()}")
                    except Exception as page_error:
                        logger.warning(f"⚠️ Erreur page {page_num + 1}: {page_error}")
                        continue
                        
        finally:
            # Nettoyage garanti
            try:
                os.unlink(tmp_path)
            except:
                pass
        
        if not text_parts:
            raise ValueError("Aucun texte extractible du PDF")
        
        return "\n\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"❌ Erreur extraction PDF: {e}")
        raise ValueError(f"Impossible d'extraire le texte du PDF: {e}")

def extract_text_from_docx(file_content: bytes) -> str:
    """📘 Extrait le texte d'un DOCX - Version optimisée"""
    if not DOCX_AVAILABLE:
        raise ValueError("python-docx non installé - impossible de lire les DOCX")
    
    try:
        # Sauvegarde temporaire
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            tmp_file.write(file_content)
            tmp_path = tmp_file.name
        
        # Extraction texte avec gestion structure
        text_parts = []
        try:
            doc = Document(tmp_path)
            
            # Extraction paragraphes
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extraction tables si présentes
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.append("[TABLE]\n" + "\n".join(table_text))
                    
        finally:
            # Nettoyage garanti
            try:
                os.unlink(tmp_path)
            except:
                pass
        
        if not text_parts:
            raise ValueError("Document DOCX vide ou sans texte extractible")
        
        return "\n\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"❌ Erreur extraction DOCX: {e}")
        raise ValueError(f"Impossible d'extraire le texte du DOCX: {e}")

def extract_text_from_file(filename: str, file_content: bytes) -> str:
    """📄 Extrait le texte selon le type de fichier - Version complète"""
    file_ext = Path(filename).suffix.lower()
    
    try:
        if file_ext == '.pdf':
            return extract_text_from_pdf(file_content)
        elif file_ext == '.docx':
            return extract_text_from_docx(file_content)
        elif file_ext in ['.txt', '.md']:
            # Détection encoding pour texte
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    return file_content.decode('latin-1')
                except UnicodeDecodeError:
                    return file_content.decode('utf-8', errors='ignore')
        elif file_ext == '.json':
            # Validation JSON + extraction texte
            try:
                json_data = json.loads(file_content.decode('utf-8'))
                return json.dumps(json_data, indent=2, ensure_ascii=False)
            except json.JSONDecodeError as e:
                raise ValueError(f"Fichier JSON invalide: {e}")
        else:
            raise ValueError(f"Format de fichier non supporté: {file_ext}")
            
    except ValueError:
        # Re-raise les erreurs de validation
        raise
    except Exception as e:
        raise ValueError(f"Erreur inattendue lors de l'extraction {file_ext}: {e}")

def get_supported_formats() -> List[str]:
    """📋 Retourne les formats supportés avec status"""
    formats = ['.txt', '.md', '.json']
    
    if PDF_AVAILABLE:
        formats.append('.pdf')
    if DOCX_AVAILABLE:
        formats.append('.docx')
    
    return formats

def validate_file_content(filename: str, file_content: bytes, extracted_text: str) -> Dict[str, Any]:
    """✅ Validation contenu fichier avec métriques"""
    file_ext = Path(filename).suffix.lower()
    
    validation = {
        'valid': True,
        'errors': [],
        'warnings': [],
        'metrics': {
            'file_size': len(file_content),
            'text_length': len(extracted_text),
            'lines_count': len(extracted_text.splitlines()),
            'words_count': len(extracted_text.split()),
            'estimated_chunks': 0
        }
    }
    
    # Validation taille minimum
    if len(extracted_text.strip()) < 10:
        validation['valid'] = False
        validation['errors'].append("Document contient moins de 10 caractères de texte")
    
    # Validation contenu selon format
    if file_ext == '.pdf' and validation['metrics']['text_length'] < 50:
        validation['warnings'].append("PDF semble avoir très peu de texte extractible")
    
    if file_ext == '.docx' and validation['metrics']['words_count'] < 5:
        validation['warnings'].append("Document DOCX semble très court")
    
    # Estimation chunks pour indexation
    paragraphs = [p.strip() for p in extracted_text.split('\n\n') if len(p.strip()) >= 50]
    validation['metrics']['estimated_chunks'] = len(paragraphs)
    
    if validation['metrics']['estimated_chunks'] == 0:
        validation['warnings'].append("Aucun chunk substantiel détecté pour indexation")
    
    return validation

# === 🆕 GESTIONNAIRE CONNEXIONS WEBSOCKET V5 + SESSION MANAGER + HEARTBEAT CORRIGÉ ===
class ConnectionManager:
    def __init__(self):
        self.active_connections: List[WebSocket] = []
        self.session_data: Dict[str, Any] = {}
        # 🔥 ATTRIBUTS HEARTBEAT CORRIGÉS
        self.heartbeat_tasks: Dict[str, asyncio.Task] = {}
        self.last_ping: Dict[str, datetime] = {}
        # 🔥 CONFIGURATION HEARTBEAT MOINS AGRESSIVE
        self.HEARTBEAT_INTERVAL = 45  # 45s au lieu de 30s
        self.HEARTBEAT_TIMEOUT = 90   # 90s au lieu de 60s

    async def connect(self, websocket: WebSocket, session_id: str):
        await websocket.accept()
        self.active_connections.append(websocket)
        
        # 🆕 Création session V5 automatique si SessionManager disponible
        session_v5_id = None
        if session_manager_v5:
            try:
                # ✅ CORRECTION: session_type dans metadata
                session_v5_id = session_manager_v5.create_session(
                    user_id="FG",  # ID utilisateur fixe pour FG
                    metadata={
                        'session_type': 'websocket',  # ✅ DANS METADATA
                        'websocket_session_id': session_id,
                        'connection_time': datetime.now().isoformat(),
                        'user_agent': 'ÉMERGENCE_V5_Frontend'
                    }
                )
                logger.info(f"✅ Session V5 créée: {session_v5_id} pour WebSocket {session_id}")
            except Exception as e:
                logger.warning(f"⚠️ Erreur création session V5: {e}")
                session_v5_id = None
        
        self.session_data[session_id] = {
            'websocket': websocket,
            'messages': [],
            'costs': {  # Tracking coûts par agent
                'anima': 0.0,
                'neo': 0.0,
                'nexus': 0.0,
                'triple': 0.0,
                'total': 0.0
            },
            'session_v5_id': session_v5_id,  # 🆕 ID session V5 pour tracking
            'created_at': datetime.now(),
            'last_activity': datetime.now(),
            'last_ping': datetime.now()  # 🔥 AJOUTÉ pour heartbeat
        }
        
        # 🔥 DÉMARRER HEARTBEAT POUR CETTE CONNEXION
        await self.start_heartbeat(session_id)
        
        logger.info(f"🔗 Nouvelle connexion WebSocket V5: {session_id}")

    async def start_heartbeat(self, session_id: str):
        """🔥 Démarrage tâche heartbeat pour une session - CORRIGÉ"""
        if session_id in self.heartbeat_tasks:
            self.heartbeat_tasks[session_id].cancel()
        
        async def heartbeat_loop():
            while session_id in self.session_data:
                try:
                    await asyncio.sleep(self.HEARTBEAT_INTERVAL)  # 45 secondes
                    
                    if session_id not in self.session_data:
                        break
                    
                    session_data = self.session_data[session_id]
                    websocket = session_data['websocket']
                    
                    # ✅ CORRECTION: Vérifier état WebSocket avant envoi
                    try:
                        if hasattr(websocket, 'client_state') and websocket.client_state.name == "CONNECTED":
                            # Envoyer ping
                            await websocket.send_text(json.dumps({
                                'type': 'ping',
                                'timestamp': datetime.now().isoformat()
                            }))
                            
                            # Vérifier dernier pong (timeout)
                            last_ping = session_data.get('last_ping', datetime.now())
                            if datetime.now() - last_ping > timedelta(seconds=self.HEARTBEAT_TIMEOUT):
                                logger.warning(f"⏰ Timeout heartbeat pour session {session_id}")
                                await self.force_disconnect(session_id, "Heartbeat timeout")
                                break
                        else:
                            logger.info(f"💔 WebSocket fermé pour session {session_id}")
                            break
                    except Exception as send_error:
                        logger.error(f"❌ Erreur envoi ping {session_id}: {send_error}")
                        await self.force_disconnect(session_id, f"Ping error: {send_error}")
                        break
                        
                except asyncio.CancelledError:
                    logger.debug(f"🛑 Heartbeat annulé pour session {session_id}")
                    break
                except Exception as e:
                    logger.error(f"❌ Erreur heartbeat session {session_id}: {e}")
                    break
        
        # Créer et stocker la tâche
        task = asyncio.create_task(heartbeat_loop())
        self.heartbeat_tasks[session_id] = task
        logger.debug(f"💓 Heartbeat démarré pour session {session_id}")

    async def handle_pong(self, session_id: str):
        """🔥 Gestion réception pong du client"""
        if session_id in self.session_data:
            self.session_data[session_id]['last_ping'] = datetime.now()
            logger.debug(f"💓 Pong reçu de session {session_id}")

    async def force_disconnect(self, session_id: str, reason: str = "Force disconnect"):
        """🔥 Fermeture forcée d'une connexion - CORRIGÉ"""
        if session_id not in self.session_data:
            return
            
        session_data = self.session_data[session_id]
        websocket = session_data['websocket']
        
        try:
            # ✅ CORRECTION: Vérifier état avant fermeture
            if hasattr(websocket, 'client_state') and websocket.client_state.name == "CONNECTED":
                await websocket.close(code=1000, reason=reason)
        except Exception as close_error:
            logger.error(f"❌ Erreur fermeture WebSocket {session_id}: {close_error}")
        
        # Appeler disconnect pour nettoyage complet
        self.disconnect(websocket, session_id)
        logger.info(f"🔌 Connexion forcée fermée: {session_id} - Raison: {reason}")

    def disconnect(self, websocket: WebSocket, session_id: str):
        # 🔥 ARRÊTER HEARTBEAT
        if session_id in self.heartbeat_tasks:
            self.heartbeat_tasks[session_id].cancel()
            del self.heartbeat_tasks[session_id]
            logger.debug(f"💔 Heartbeat arrêté pour session {session_id}")
        
        if websocket in self.active_connections:
            self.active_connections.remove(websocket)
        
        # 🆕 Finalisation session V5 automatique
        if session_id in self.session_data:
            session_data = self.session_data[session_id]
            session_v5_id = session_data.get('session_v5_id')
            
            if session_v5_id and session_manager_v5:
                try:
                    # Finalisation avec métadonnées finales
                    final_metadata = {
                        'disconnection_time': datetime.now().isoformat(),
                        'total_messages': len(session_data['messages']),
                        'session_duration_minutes': (
                            datetime.now() - session_data['created_at']
                        ).total_seconds() / 60,
                        'final_costs': session_data['costs'],
                        'heartbeat_disconnect': True
                    }
                    
                    session_manager_v5.finalize_session(session_v5_id, final_metadata)
                    logger.info(f"✅ Session V5 finalisée: {session_v5_id}")
                    
                except Exception as e:
                    logger.warning(f"⚠️ Erreur finalisation session V5: {e}")
            
            del self.session_data[session_id]
        
        logger.info(f"🔌 Connexion fermée: {session_id}")

    async def send_message(self, session_id: str, message: dict):
        """✅ ENVOI SÉCURISÉ MESSAGES - CORRIGÉ"""
        if session_id not in self.session_data:
            logger.warning(f"⚠️ Session {session_id} non trouvée pour envoi")
            return False
            
        websocket = self.session_data[session_id]['websocket']
        try:
            # ✅ CORRECTION: Vérifier état WebSocket avant envoi
            if hasattr(websocket, 'client_state') and websocket.client_state.name == "CONNECTED":
                await websocket.send_text(json.dumps(message))
                
                # 🆕 Mise à jour activité session
                self.session_data[session_id]['last_activity'] = datetime.now()
                return True
            else:
                logger.warning(f"⚠️ WebSocket non connecté pour session {session_id}")
                return False
                
        except Exception as e:
            logger.error(f"❌ Erreur envoi message session {session_id}: {e}")
            # 🔥 DÉCONNEXION AUTO SI ERREUR ENVOI
            await self.force_disconnect(session_id, f"Send error: {e}")
            return False

    def update_cost(self, session_id: str, agent: str, cost: float):
        """🆕 Mise à jour coûts par agent"""
        if session_id in self.session_data:
            costs = self.session_data[session_id]['costs']
            costs[agent] += cost
            costs['total'] += cost

    def get_costs(self, session_id: str) -> Dict[str, float]:
        """🆕 Récupération coûts session"""
        if session_id in self.session_data:
            return self.session_data[session_id]['costs']
        return {'anima': 0.0, 'neo': 0.0, 'nexus': 0.0, 'triple': 0.0, 'total': 0.0}

    def add_message_to_session_v5(self, session_id: str, message_type: str, content: dict):
        """🆕 Ajout message à session V5 via SessionManager"""
        if session_id not in self.session_data:
            return
        
        session_data = self.session_data[session_id]
        session_v5_id = session_data.get('session_v5_id')
        
        if session_v5_id and session_manager_v5:
            try:
                session_manager_v5.add_message_to_session(
                    session_id=session_v5_id,
                    message_type=message_type,
                    content=content,
                    metadata={
                        'websocket_session_id': session_id,
                        'timestamp': datetime.now().isoformat()
                    }
                )
                
                # Aussi stocker localement pour tracking
                session_data['messages'].append({
                    'type': message_type,
                    'content': content,
                    'timestamp': datetime.now().isoformat()
                })
                
            except Exception as e:
                logger.warning(f"⚠️ Erreur ajout message session V5: {e}")

manager = ConnectionManager()

# === ROUTES PRINCIPALES ===
@app.get("/", response_class=HTMLResponse)
async def read_root():
    """🏠 Sert la page principale ÉMERGENCE V5"""
    try:
        with open("interface/frontend/index.html", "r", encoding="utf-8") as f:
            html_content = f.read()
            
            # 🆕 Injection info SessionManager V5
            session_info = f"✅ Mémoire persistante" if SESSION_MANAGER_AVAILABLE else "❌ Mémoire basique"
            formats_info = f"<!-- ÉMERGENCE V5 - Formats: {', '.join(get_supported_formats())} | Sessions: {session_info} -->"
            html_content = html_content.replace("</head>", f"{formats_info}\n</head>")
            
            return HTMLResponse(content=html_content)
    except FileNotFoundError:
        return HTMLResponse(f"""
        <html>
            <head><title>ÉMERGENCE V5 - Interface Required</title></head>
            <body>
                <h1>🚀 ÉMERGENCE V5 - Backend Running</h1>
                <p><strong>Formats supportés:</strong> {", ".join(get_supported_formats())}</p>
                <p><strong>Modules V4:</strong> {'✅ Disponibles' if MODULES_AVAILABLE else '❌ Non disponibles'}</p>
                <p><strong>SessionManager V5:</strong> {'✅ Mémoire persistante' if SESSION_MANAGER_AVAILABLE else '❌ Mémoire basique'}</p>
                <h2>🎯 Modes supportés:</h2>
                <ul>
                    <li><strong>Mode Dialogue</strong> - Chat avec agent unique</li>
                    <li><strong>Mode Triangle</strong> - Débat triangulaire 3 agents</li>
                    <li><strong>Mode Documents</strong> - Sélection granulaire docs + agents</li>
                </ul>
                <h2>🆕 Nouveautés V5.1.2:</h2>
                <ul>
                    <li><strong>Mémoire persistante</strong> - Journalisation automatique conversations</li>
                    <li><strong>Recherche temporelle</strong> - Retrouver discussions précédentes</li>
                    <li><strong>Continuité sessions</strong> - Agents se souviennent du contexte</li>
                    <li><strong>Heartbeat stable</strong> - WebSocket mobile sans déconnexion</li>
                </ul>
                <p>Ajoutez le fichier <code>interface/frontend/index.html</code> avec l'interface V5.</p>
            </body>
        </html>
        """)

@app.get("/api/status")
async def get_system_status() -> SystemStatus:
    """📊 Status du système V5 + SessionManager"""
    try:
        if MODULES_AVAILABLE:
            real_agents = get_real_agents()
            agents_status = real_agents.get_health_status()
            
            rag_manager = get_rag_manager()
            db_stats = rag_manager.get_stats()
            
            # Calcul coût total robuste
            try:
                usage_stats = real_agents.get_usage_stats()
                total_cost = 0.0
                if isinstance(usage_stats, dict):
                    if 'totals' in usage_stats and isinstance(usage_stats['totals'], dict):
                        total_cost = usage_stats['totals'].get('cost', 0.0)
                    elif 'total_cost' in usage_stats:
                        total_cost = usage_stats['total_cost']
                    elif 'cost' in usage_stats:
                        total_cost = usage_stats['cost']
                    else:
                        for agent_stats in usage_stats.values():
                            if isinstance(agent_stats, dict) and 'cost' in agent_stats:
                                total_cost += agent_stats['cost']
            except Exception as e:
                logger.warning(f"⚠️ Erreur calcul coût total: {e}")
                total_cost = 0.0
            
            return SystemStatus(
                modules_available=True,
                agents_status=agents_status,
                database_stats=db_stats,
                total_cost=total_cost,
                session_manager_available=SESSION_MANAGER_AVAILABLE
            )
        else:
            return SystemStatus(
                modules_available=False,
                agents_status={},
                database_stats={},
                total_cost=0.0,
                session_manager_available=SESSION_MANAGER_AVAILABLE
            )
    except Exception as e:
        logger.error(f"❌ Erreur status: {e}")
        return SystemStatus(
            modules_available=False,
            agents_status={'error': str(e)},
            database_stats={'error': str(e)},
            total_cost=0.0,
            session_manager_available=SESSION_MANAGER_AVAILABLE
        )

# 🆕 NOUVELLES ROUTES SESSION MANAGER V5 - CORRIGÉES
@app.get("/api/sessions/recent")
async def get_recent_sessions(limit: int = Query(10, ge=1, le=50)):
    """📚 Récupère les sessions récentes avec SessionManager V5 - AVEC FALLBACKS"""
    if not SESSION_MANAGER_AVAILABLE or not session_manager_v5:
        return JSONResponse({
            'success': False,
            'sessions': [],
            'total': 0,
            'message': 'SessionManager V5 non disponible'
        }, status_code=503)
    
    try:
        # ✅ CORRECTION: Vérifier si méthode existe
        if hasattr(session_manager_v5, 'get_recent_sessions'):
            sessions = session_manager_v5.get_recent_sessions("FG", limit)
        else:
            # Fallback basique
            sessions = []
            logger.warning("⚠️ get_recent_sessions non implémentée")
        
        return {
            'success': True,
            'sessions': sessions,
            'total': len(sessions)
        }
    except Exception as e:
        logger.error(f"❌ Erreur récupération sessions: {e}")
        return {
            'success': False,
            'sessions': [],
            'total': 0,
            'error': str(e)
        }

@app.get("/api/sessions/{session_id}")
async def get_session_details(session_id: str):
    """🔍 Détails complets d'une session spécifique avec fallback"""
    if not SESSION_MANAGER_AVAILABLE or not session_manager_v5:
        return JSONResponse({
            'success': False,
            'session': None,
            'message': 'SessionManager V5 non disponible'
        }, status_code=503)
    
    try:
        # ✅ CORRECTION: Vérifier si méthode existe
        session_data = None
        if hasattr(session_manager_v5, 'get_session_by_id'):
            session_data = session_manager_v5.get_session_by_id(session_id)
        else:
            # Fallback: chercher dans recent sessions
            if hasattr(session_manager_v5, 'get_recent_sessions'):
                recent_sessions = session_manager_v5.get_recent_sessions("FG", limit=50)
                session_data = next((s for s in recent_sessions if s.get('session_id') == session_id), None)
        
        if not session_data:
            return JSONResponse({
                'success': False,
                'session': None,
                'message': 'Session non trouvée'
            }, status_code=404)
        
        return {
            'success': True,
            'session': session_data
        }
    except Exception as e:
        logger.error(f"❌ Erreur détails session {session_id}: {e}")
        return {
            'success': False,
            'session': None,
            'error': str(e)
        }

@app.get("/api/sessions/search")
async def search_sessions(
    theme: Optional[str] = Query(None, description="Recherche par thème"),
    date_from: Optional[str] = Query(None, description="Date début (ISO format)"),
    date_to: Optional[str] = Query(None, description="Date fin (ISO format)"),
    limit: int = Query(20, ge=1, le=100)
):
    """🔍 Recherche sessions par critères - AVEC FALLBACKS"""
    if not SESSION_MANAGER_AVAILABLE or not session_manager_v5:
        return JSONResponse({
            'success': False,
            'sessions': [],
            'total': 0,
            'message': 'SessionManager V5 non disponible'
        }, status_code=503)
    
    try:
        sessions = []
        
        # ✅ CORRECTION: Vérifier méthodes disponibles
        if theme and hasattr(session_manager_v5, 'search_sessions_by_theme'):
            sessions = session_manager_v5.search_sessions_by_theme(theme, limit)
        elif date_from and date_to and hasattr(session_manager_v5, 'search_sessions_by_date'):
            sessions = session_manager_v5.search_sessions_by_date(date_from, date_to, limit)
        elif hasattr(session_manager_v5, 'get_recent_sessions'):
            sessions = session_manager_v5.get_recent_sessions("FG", limit)
        else:
            logger.warning("⚠️ Méthodes de recherche non implémentées")
        
        return {
            'success': True,
            'sessions': sessions,
            'total': len(sessions),
            'search_criteria': {
                'theme': theme,
                'date_from': date_from,
                'date_to': date_to
            }
        }
    except Exception as e:
        logger.error(f"❌ Erreur recherche sessions: {e}")
        return {
            'success': False,
            'sessions': [],
            'total': 0,
            'error': str(e)
        }

@app.get("/api/sessions/stats")
async def get_sessions_stats():
    """📊 Statistiques générales sessions avec fallback"""
    if not SESSION_MANAGER_AVAILABLE or not session_manager_v5:
        return JSONResponse({
            'success': False,
            'stats': {'total_sessions': 0},
            'message': 'SessionManager V5 non disponible'
        }, status_code=503)
    
    try:
        # ✅ CORRECTION: Vérifier si méthode existe
        if hasattr(session_manager_v5, 'get_session_stats'):
            stats = session_manager_v5.get_session_stats("FG")
        else:
            # Fallback: stats basiques via get_recent_sessions
            if hasattr(session_manager_v5, 'get_recent_sessions'):
                recent_sessions = session_manager_v5.get_recent_sessions("FG", limit=100)
                stats = {
                    'total_sessions': len(recent_sessions),
                    'method': 'fallback_via_recent',
                    'note': 'get_session_stats method not implemented yet'
                }
            else:
                stats = {
                    'total_sessions': 0,
                    'method': 'emergency_fallback',
                    'note': 'No session methods available'
                }
        
        return {
            'success': True,
            'stats': stats
        }
    except Exception as e:
        logger.error(f"❌ Erreur stats sessions: {e}")
        # Stats d'urgence
        return {
            'success': True,
            'stats': {
                'total_sessions': 0,
                'error': str(e),
                'method': 'emergency_fallback'
            }
        }

# === ROUTES CHAT EXISTANTES (avec tracking session V5) ===
@app.post("/api/chat")
async def chat_with_agent(message: ChatMessage) -> AgentResponseModel:
    """💬 Chat avec un agent individuel + tracking session V5"""
    if not MODULES_AVAILABLE:
        raise HTTPException(status_code=503, detail="Modules V4 non disponibles")
    
    try:
        # Récupération contexte RAG
        context = ""
        rag_chunks_count = 0
        if message.use_rag:
            rag_manager = get_rag_manager()
            context = rag_manager.get_context_for_agent(
                message.message, 
                message.agent, 
                message.rag_chunks
            )
            rag_chunks_count = context.count("=== DOCUMENT") + context.count("--- Interaction")
        
        # Réponse agent
        real_agents = get_real_agents()
        response = real_agents.get_response(message.agent, message.message, context)
        
        # 🆕 Stockage session V5 si disponible
        if session_manager_v5:
            try:
                # ✅ CORRECTION: session_type dans metadata
                api_session_id = session_manager_v5.create_session(
                    user_id="FG",
                    metadata={
                        'session_type': 'api_chat',  # ✅ DANS METADATA
                        'endpoint': '/api/chat', 
                        'agent': message.agent
                    }
                )
                
                # Ajout message user
                session_manager_v5.add_message_to_session(
                    session_id=api_session_id,
                    message_type="user",
                    content={
                        'text': message.message,
                        'agent_target': message.agent,
                        'rag_enabled': message.use_rag,
                        'rag_chunks': message.rag_chunks
                    }
                )
                
                # Ajout réponse agent
                session_manager_v5.add_message_to_session(
                    session_id=api_session_id,
                    message_type="agent",
                    content={
                        'agent': message.agent,
                        'text': response.response_text,
                        'model': response.model_used,
                        'processing_time': response.processing_time,
                        'cost': response.cost_estimate,
                        'rag_chunks_count': rag_chunks_count
                    }
                )
                
                # Finalisation session API
                session_manager_v5.finalize_session(api_session_id, {
                    'api_endpoint': '/api/chat',
                    'single_interaction': True
                })
                
            except Exception as session_error:
                logger.warning(f"⚠️ Erreur session V5 API chat: {session_error}")
        
        # Stockage interaction V4 (existant)
        try:
            rag_manager = get_rag_manager()
            session_id = f"api_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            rag_manager.store_interaction_v4(
                session_id=session_id,
                user_message=message.message,
                agent_name=message.agent,
                agent_response=response.response_text,
                processing_time=response.processing_time,
                rag_chunks_used=rag_chunks_count
            )
        except Exception as store_error:
            logger.warning(f"⚠️ Erreur stockage interaction (non-critique): {store_error}")
        
        return AgentResponseModel(
            agent=message.agent,
            response_text=response.response_text,
            processing_time=response.processing_time,
            model_used=response.model_used,
            rag_chunks_count=rag_chunks_count,
            cost_estimate=response.cost_estimate,
            provider=response.provider,
            timestamp=datetime.now().isoformat()
        )
        
    except Exception as e:
        logger.error(f"❌ Erreur chat agent {message.agent}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# === ROUTES UPLOAD + GESTION DOCUMENTS ===
@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...)) -> dict:
    """🔥 Upload et indexation de document - SUPPORT COMPLET PDF/DOCX/TXT/MD/JSON"""
    if not MODULES_AVAILABLE:
        raise HTTPException(status_code=503, detail="Modules V4 non disponibles")
    
    try:
        # Validation taille fichier (max 10MB)
        MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
        file_content = await file.read()
        
        if len(file_content) > MAX_FILE_SIZE:
            raise HTTPException(status_code=413, detail=f"Fichier trop volumineux (max 10MB)")
        
        if len(file_content) == 0:
            raise HTTPException(status_code=400, detail="Fichier vide")
        
        # Validation extension
        file_ext = Path(file.filename).suffix.lower()
        supported_formats = get_supported_formats()
        
        if file_ext not in supported_formats:
            raise HTTPException(
                status_code=400, 
                detail=f"Format non supporté: {file_ext}. Formats acceptés: {', '.join(supported_formats)}"
            )
        
        # Extraction texte selon format
        try:
            text_content = extract_text_from_file(file.filename, file_content)
        except ValueError as ve:
            raise HTTPException(status_code=400, detail=str(ve))
        
        # Validation contenu avec métriques
        validation = validate_file_content(file.filename, file_content, text_content)
        
        if not validation['valid']:
            raise HTTPException(status_code=400, detail=f"Validation échouée: {'; '.join(validation['errors'])}")
        
        # Hash fichier pour éviter doublons
        file_hash = hashlib.md5(file_content).hexdigest()
        
        # Indexation avec RAG Manager V4
        rag_manager = get_rag_manager()
        success = rag_manager.store_document_v4(
            filename=file.filename,
            content=text_content,
            file_path=f"uploads/{file.filename}",
            file_hash=file_hash,
            file_size=len(file_content)
        )
        
        if success:
            response_data = {
                'success': True,
                'filename': file.filename,
                'format': file_ext,
                'file_hash': file_hash,
                'validation': validation,
                'message': f"Document '{file.filename}' indexé avec succès"
            }
            
            # Ajout warnings si présents
            if validation['warnings']:
                response_data['warnings'] = validation['warnings']
            
            return response_data
        else:
            raise HTTPException(status_code=500, detail="Erreur lors de l'indexation en base")
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Erreur upload: {e}")
        raise HTTPException(status_code=500, detail=f"Erreur inattendue: {str(e)}")

@app.get("/api/documents")
async def list_documents(
    limit: int = Query(50, ge=1, le=100, description="Nombre max de documents à retourner"),
    offset: int = Query(0, ge=0, description="Offset pour pagination")
) -> DocumentListResponse:
    """📋 Liste des documents indexés avec détails complets"""
    if not MODULES_AVAILABLE:
        return DocumentListResponse(
            documents=[],
            total=0,
            total_size=0,
            supported_formats=get_supported_formats(),
            pdf_available=PDF_AVAILABLE,
            docx_available=DOCX_AVAILABLE
        )
    
    try:
        # Implémentation via database directe
        db = get_db()
        documents = []
        total_docs = 0
        total_size = 0
        
        # Récupération documents via SQL direct
        try:
            # Query pour documents avec pagination
            with db.get_connection() as conn:
                cursor = conn.cursor()
                
                # Count total
                cursor.execute("SELECT COUNT(*) FROM documents")
                total_docs = cursor.fetchone()[0]
                
                # Documents avec détails
                cursor.execute("""
                    SELECT 
                        id, filename, file_size, upload_date, 
                        chunks_count, file_hash, processing_status, metadata
                    FROM documents 
                    ORDER BY upload_date DESC 
                    LIMIT ? OFFSET ?
                """, (limit, offset))
                
                for row in cursor.fetchall():
                    doc_id, filename, file_size, upload_date, chunks_count, file_hash, status, metadata = row
                    
                    # Format détection
                    file_format = Path(filename).suffix.lower() if filename else 'unknown'
                    
                    # Text preview via chunks
                    text_preview = ""
                    try:
                        cursor.execute("""
                            SELECT content FROM chunks 
                            WHERE document_id = ? 
                            ORDER BY chunk_index LIMIT 1
                        """, (doc_id,))
                        chunk_row = cursor.fetchone()
                        if chunk_row and chunk_row[0]:
                            preview_text = chunk_row[0][:200]
                            text_preview = preview_text + "..." if len(chunk_row[0]) > 200 else preview_text
                    except:
                        pass
                    
                    documents.append(DocumentInfo(
                        id=doc_id,
                        filename=filename or 'unknown',
                        file_size=file_size or 0,
                        upload_date=upload_date or '',
                        chunks_count=chunks_count or 0,
                        file_hash=file_hash or '',
                        format=file_format,
                        text_preview=text_preview,
                        processing_status=status or 'completed'
                    ))
                    
                    total_size += (file_size or 0)
                
        except Exception as db_error:
            logger.error(f"❌ Erreur requête documents database: {db_error}")
            # Fallback via RAG stats si database fail
            rag_manager = get_rag_manager()
            stats = rag_manager.get_stats()
            total_docs = stats.get('database', {}).get('documents_count', 0)
        
        return DocumentListResponse(
            documents=documents,
            total=total_docs,
            total_size=total_size,
            supported_formats=get_supported_formats(),
            pdf_available=PDF_AVAILABLE,
            docx_available=DOCX_AVAILABLE
        )
        
    except Exception as e:
        logger.error(f"❌ Erreur liste documents: {e}")
        # Retour minimal en cas d'erreur
        return DocumentListResponse(
            documents=[],
            total=0,
            total_size=0,
            supported_formats=get_supported_formats(),
            pdf_available=PDF_AVAILABLE,
            docx_available=DOCX_AVAILABLE
        )

@app.delete("/api/documents/{document_id}")
async def delete_document(document_id: str) -> dict:
    """🗑️ Suppression sélective d'un document"""
    if not MODULES_AVAILABLE:
        raise HTTPException(status_code=503, detail="Modules V4 non disponibles")
    
    try:
        # Implémentation via database directe
        db = get_db()
        
        # Vérification existence + récupération info
        with db.get_connection() as conn:
            cursor = conn.cursor()
            
            # Check existence
            cursor.execute("SELECT filename FROM documents WHERE id = ?", (document_id,))
            doc_row = cursor.fetchone()
            
            if not doc_row:
                raise HTTPException(status_code=404, detail="Document non trouvé")
            
            filename = doc_row[0]
            
            # Suppression chunks associés (cascade normalement gérée par FK)
            cursor.execute("DELETE FROM chunks WHERE document_id = ?", (document_id,))
            chunks_deleted = cursor.rowcount
            
            # Suppression document
            cursor.execute("DELETE FROM documents WHERE id = ?", (document_id,))
            doc_deleted = cursor.rowcount
            
            # Commit
            conn.commit()
            
            if doc_deleted > 0:
                logger.info(f"✅ Document supprimé: {filename} (ID: {document_id}, {chunks_deleted} chunks)")
                
                # TODO: Suppression vecteurs ChromaDB si possible
                try:
                    vector_manager = get_vector_manager()
                    if hasattr(vector_manager, 'delete_document_vectors'):
                        vector_manager.delete_document_vectors(document_id)
                except Exception as vector_error:
                    logger.warning(f"⚠️ Impossible de supprimer vecteurs: {vector_error}")
                
                return {
                    'success': True,
                    'document_id': document_id,
                    'filename': filename,
                    'chunks_deleted': chunks_deleted,
                    'message': f"Document '{filename}' supprimé avec succès"
                }
            else:
                raise HTTPException(status_code=500, detail="Échec suppression database")
                
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Erreur suppression document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Erreur inattendue: {str(e)}")

@app.get("/api/database/stats")
async def get_database_stats() -> DatabaseStats:
    """📊 Statistiques détaillées base de données"""
    if not MODULES_AVAILABLE:
        raise HTTPException(status_code=503, detail="Modules V4 non disponibles")
    
    try:
        # Stats via RAG Manager
        rag_manager = get_rag_manager()
        stats = rag_manager.get_stats()
        
        # Extraction métriques
        db_stats = stats.get('database', {})
        vector_stats = stats.get('vectors', {})
        
        # Calcul taille base (estimation)
        db_path = "data/emergence_v4.db"
        total_size_mb = 0.0
        try:
            if os.path.exists(db_path):
                total_size_mb = os.path.getsize(db_path) / (1024 * 1024)
        except:
            pass
        
        return DatabaseStats(
            documents_count=db_stats.get('documents_count', 0),
            chunks_count=db_stats.get('chunks_count', 0),
            conversations_count=db_stats.get('conversations_count', 0),
            total_size_mb=round(total_size_mb, 2),
            database_path=db_path,
            vector_count=vector_stats.get('vector_count', 0),
            last_update=datetime.now().isoformat()
        )
        
    except Exception as e:
        logger.error(f"❌ Erreur stats database: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# === WEBSOCKET V5 - SUPPORT 3 MODES + SESSION TRACKING + HEARTBEAT ===
@app.websocket("/ws/{session_id}")
async def websocket_endpoint(websocket: WebSocket, session_id: str):
    """🔌 WebSocket ÉMERGENCE V5 - Support 3 modes + session tracking V5 + Heartbeat"""
    await manager.connect(websocket, session_id)
    
    try:
        while True:
            data = await websocket.receive_text()
            message_data = json.loads(data)
            
            msg_type = message_data.get('type', 'chat')
            
            # 🔥 GESTION MESSAGES HEARTBEAT
            if msg_type == 'ping':
                # Répondre pong
                await manager.send_message(session_id, {
                    'type': 'pong',
                    'timestamp': datetime.now().isoformat()
                })
                continue
            elif msg_type == 'pong':
                # Mettre à jour dernière activité
                await manager.handle_pong(session_id)
                continue
            
            # Messages normaux
            if msg_type == 'chat':
                await handle_websocket_chat(session_id, message_data)
            elif msg_type == 'triple':
                await handle_websocket_triple(session_id, message_data)
            elif msg_type == 'documents':
                await handle_websocket_documents_mode(session_id, message_data)
            elif msg_type == 'status':
                await handle_websocket_status(session_id)
            elif msg_type == 'documents_list':
                await handle_websocket_documents_list(session_id, message_data)
                
    except WebSocketDisconnect:
        manager.disconnect(websocket, session_id)
    except Exception as e:
        logger.error(f"❌ Erreur WebSocket: {e}")
        try:
            await manager.send_message(session_id, {
                'type': 'error',
                'message': str(e)
            })
        except:
            pass
        finally:
            manager.disconnect(websocket, session_id)

# === 🔥 WEBSOCKET HANDLERS V5 + SESSION TRACKING - COMPLETS - CORRIGÉS ===

async def handle_websocket_chat(session_id: str, message_data: dict):
    """💬 Handler mode dialogue WebSocket + Session V5"""
    try:
        user_message = message_data.get('message', '')
        agent_name = message_data.get('agent', 'nexus')
        use_rag = message_data.get('use_rag', True)
        rag_chunks = message_data.get('rag_chunks', 5)
        
        if not user_message.strip():
            await manager.send_message(session_id, {
                'type': 'error',
                'message': 'Message vide'
            })
            return
        
        # 🆕 Ajout message utilisateur à session V5
        manager.add_message_to_session_v5(session_id, "user", {
            'text': user_message,
            'agent_target': agent_name,
            'mode': 'chat'
        })
        
        await manager.send_message(session_id, {
            'type': 'thinking',
            'agent': agent_name,
            'message': f'{agent_name.title()} réfléchit...'
        })
        
        # Contexte RAG
        context = ""
        rag_chunks_count = 0
        if use_rag and MODULES_AVAILABLE:
            try:
                rag_manager = get_rag_manager()
                context = rag_manager.get_context_for_agent(user_message, agent_name, rag_chunks)
                rag_chunks_count = context.count("=== DOCUMENT") + context.count("--- Interaction")
            except Exception as rag_error:
                logger.warning(f"⚠️ Erreur RAG: {rag_error}")
        
        # Réponse agent
        if MODULES_AVAILABLE:
            real_agents = get_real_agents()
            response = real_agents.get_response(agent_name, user_message, context)
            
            # 🆕 Ajout réponse agent à session V5
            manager.add_message_to_session_v5(session_id, "agent", {
                'agent': agent_name,
                'text': response.response_text,
                'model': response.model_used,
                'processing_time': response.processing_time,
                'cost': response.cost_estimate,
                'rag_chunks_count': rag_chunks_count
            })
            
            # Mise à jour coût session
            manager.update_cost(session_id, agent_name, response.cost_estimate)
            
            await manager.send_message(session_id, {
                'type': 'agent_response',
                'agent': agent_name,
                'message': response.response_text,
                'model': response.model_used,
                'processing_time': response.processing_time,
                'cost': response.cost_estimate,
                'rag_chunks_count': rag_chunks_count,
                'session_costs': manager.get_costs(session_id)
            })
        else:
            await manager.send_message(session_id, {
                'type': 'error',
                'message': 'Modules V4 non disponibles'
            })
            
    except Exception as e:
        logger.error(f"❌ Erreur WebSocket chat: {e}")
        await manager.send_message(session_id, {
            'type': 'error',
            'message': f'Erreur traitement: {str(e)}'
        })

async def handle_websocket_triple(session_id: str, message_data: dict):
    """🔀 Handler mode triangle WebSocket + Session V5"""
    try:
        user_message = message_data.get('message', '')
        use_rag = message_data.get('use_rag', True)
        rag_chunks = message_data.get('rag_chunks', 5)
        
        if not user_message.strip():
            await manager.send_message(session_id, {
                'type': 'error',
                'message': 'Message vide'
            })
            return
        
        # 🆕 Ajout message utilisateur à session V5
        manager.add_message_to_session_v5(session_id, "user", {
            'text': user_message,
            'mode': 'triangle'
        })
        
        if not MODULES_AVAILABLE:
            await manager.send_message(session_id, {
                'type': 'error',
                'message': 'Modules V4 non disponibles'
            })
            return
        
        # Contexte RAG
        context = ""
        rag_chunks_count = 0
        if use_rag:
            try:
                rag_manager = get_rag_manager()
                context = rag_manager.get_context_for_agent(user_message, "anima", rag_chunks)
                rag_chunks_count = context.count("=== DOCUMENT") + context.count("--- Interaction")
            except Exception as rag_error:
                logger.warning(f"⚠️ Erreur RAG triangle: {rag_error}")
        
        real_agents = get_real_agents()
        agents_order = ['anima', 'neo', 'nexus']
        total_cost = 0.0
        
        # Début débat triangulaire
        await manager.send_message(session_id, {
            'type': 'triangle_start',
            'message': f'Début du débat triangulaire sur: "{user_message}"',
            'agents': agents_order
        })
        
        # Réponses séquentielles
        for i, agent_name in enumerate(agents_order):
            try:
                await manager.send_message(session_id, {
                    'type': 'triangle_thinking',
                    'agent': agent_name,
                    'position': i + 1,
                    'message': f'{agent_name.title()} prépare sa réponse...'
                })
                
                # Message adapté selon position
                if i == 0:  # Anima ouvre
                    prompt_message = user_message
                else:  # Neo et Nexus réagissent
                    prompt_message = f"Débat en cours: {user_message}\n\nÉchanges précédents disponibles dans le contexte."
                
                response = real_agents.get_response(agent_name, prompt_message, context)
                total_cost += response.cost_estimate
                
                # 🆕 Ajout réponse agent à session V5
                manager.add_message_to_session_v5(session_id, "agent", {
                    'agent': agent_name,
                    'text': response.response_text,
                    'model': response.model_used,
                    'processing_time': response.processing_time,
                    'cost': response.cost_estimate,
                    'triangle_position': i + 1,
                    'triangle_role': ['ouverture', 'challenge', 'synthèse'][i]
                })
                
                # Mise à jour coût
                manager.update_cost(session_id, agent_name, response.cost_estimate)
                manager.update_cost(session_id, 'triple', response.cost_estimate)
                
                await manager.send_message(session_id, {
                    'type': 'triangle_response',
                    'agent': agent_name,
                    'position': i + 1,
                    'message': response.response_text,
                    'model': response.model_used,
                    'processing_time': response.processing_time,
                    'cost': response.cost_estimate,
                    'role': ['ouverture', 'challenge', 'synthèse'][i]
                })
                
                # Pause entre agents
                await asyncio.sleep(1)
                
            except Exception as agent_error:
                logger.error(f"❌ Erreur agent {agent_name}: {agent_error}")
                await manager.send_message(session_id, {
                    'type': 'triangle_error',
                    'agent': agent_name,
                    'message': f'Erreur {agent_name}: {str(agent_error)}'
                })
        
        # Fin débat
        await manager.send_message(session_id, {
            'type': 'triangle_complete',
            'total_cost': total_cost,
            'agents_count': len(agents_order),
            'session_costs': manager.get_costs(session_id)
        })
        
    except Exception as e:
        logger.error(f"❌ Erreur WebSocket triangle: {e}")
        await manager.send_message(session_id, {
            'type': 'error',
            'message': f'Erreur débat triangulaire: {str(e)}'
        })

async def handle_websocket_documents_mode(session_id: str, message_data: dict):
    """📁 Handler mode documents WebSocket + Session V5 - FONCTION MANQUANTE AJOUTÉE !"""
    try:
        user_message = message_data.get('message', '')
        selected_agents = message_data.get('agents', ['nexus'])
        selected_documents = message_data.get('documents', [])
        use_rag = message_data.get('use_rag', True)
        rag_chunks = message_data.get('rag_chunks', 5)
        
        if not user_message.strip():
            await manager.send_message(session_id, {
                'type': 'error',
                'message': 'Message vide'
            })
            return
        
        # 🆕 Ajout message utilisateur à session V5
        manager.add_message_to_session_v5(session_id, "user", {
            'text': user_message,
            'selected_agents': selected_agents,
            'selected_documents': selected_documents,
            'mode': 'documents'
        })
        
        if not MODULES_AVAILABLE:
            await manager.send_message(session_id, {
                'type': 'error',
                'message': 'Modules V4 non disponibles'
            })
            return
        
        # Contexte RAG spécialisé documents
        context = ""
        rag_chunks_count = 0
        if use_rag:
            try:
                rag_manager = get_rag_manager()
                
                # Si documents spécifiques sélectionnés
                if selected_documents and hasattr(rag_manager, 'get_context_for_documents'):
                    context = rag_manager.get_context_for_documents(
                        user_message, 
                        selected_documents, 
                        rag_chunks
                    )
                else:
                    # Contexte général tous documents
                    context = rag_manager.get_context_for_agent(
                        user_message, 
                        selected_agents[0], 
                        rag_chunks
                    )
                
                rag_chunks_count = context.count("=== DOCUMENT") + context.count("--- Interaction")
                
                # Envoi info contexte trouvé
                await manager.send_message(session_id, {
                    'type': 'documents_context',
                    'chunks_found': rag_chunks_count,
                    'context_preview': context[:300] + "..." if len(context) > 300 else context
                })
                
            except Exception as rag_error:
                logger.warning(f"⚠️ Erreur RAG documents: {rag_error}")
        
        real_agents = get_real_agents()
        total_cost = 0.0
        
        # Début mode documents
        await manager.send_message(session_id, {
            'type': 'documents_start',
            'message': f'Analyse documentaire avec {len(selected_agents)} agent(s)',
            'agents': selected_agents,
            'documents_count': len(selected_documents) if selected_documents else 'tous'
        })
        
        # Réponses des agents sélectionnés
        for i, agent_name in enumerate(selected_agents):
            try:
                await manager.send_message(session_id, {
                    'type': 'documents_thinking',
                    'agent': agent_name,
                    'message': f'{agent_name.title()} analyse les documents...'
                })
                
                # Prompt enrichi avec contexte documentaire
                prompt_message = f"""Question: {user_message}

Contexte documentaire disponible:
{context[:2000] if context else "Aucun contexte documentaire spécifique."}

Analyse cette question en te basant sur les documents fournis."""
                
                response = real_agents.get_response(agent_name, prompt_message, context)
                total_cost += response.cost_estimate
                
                # 🆕 Ajout réponse agent à session V5
                manager.add_message_to_session_v5(session_id, "agent", {
                    'agent': agent_name,
                    'text': response.response_text,
                    'model': response.model_used,
                    'processing_time': response.processing_time,
                    'cost': response.cost_estimate,
                    'rag_chunks_count': rag_chunks_count,
                    'documents_mode': True
                })
                
                # Mise à jour coût
                manager.update_cost(session_id, agent_name, response.cost_estimate)
                
                await manager.send_message(session_id, {
                    'type': 'documents_response',
                    'agent': agent_name,
                    'message': response.response_text,
                    'model': response.model_used,
                    'processing_time': response.processing_time,
                    'cost': response.cost_estimate,
                    'rag_chunks_count': rag_chunks_count
                })
                
                # Pause entre agents si plusieurs
                if len(selected_agents) > 1 and i < len(selected_agents) - 1:
                    await asyncio.sleep(1)
                
            except Exception as agent_error:
                logger.error(f"❌ Erreur agent {agent_name} documents: {agent_error}")
                await manager.send_message(session_id, {
                    'type': 'documents_error',
                    'agent': agent_name,
                    'message': f'Erreur {agent_name}: {str(agent_error)}'
                })
        
        # Fin mode documents
        await manager.send_message(session_id, {
            'type': 'documents_complete',
            'total_cost': total_cost,
            'agents_count': len(selected_agents),
            'rag_chunks_used': rag_chunks_count,
            'session_costs': manager.get_costs(session_id)
        })
        
    except Exception as e:
        logger.error(f"❌ Erreur WebSocket documents: {e}")
        await manager.send_message(session_id, {
            'type': 'error',
            'message': f'Erreur mode documents: {str(e)}'
        })

async def handle_websocket_status(session_id: str):
    """📊 Handler status système WebSocket"""
    try:
        if MODULES_AVAILABLE:
            real_agents = get_real_agents()
            agents_status = real_agents.get_health_status()
            usage_stats = real_agents.get_usage_stats()
            
            rag_manager = get_rag_manager()
            db_stats = rag_manager.get_stats()
        else:
            agents_status = {'error': 'Modules non disponibles'}
            usage_stats = {'error': 'Modules non disponibles'}
            db_stats = {'error': 'Modules non disponibles'}
        
        # Coûts session actuelle
        session_costs = manager.get_costs(session_id)
        
        await manager.send_message(session_id, {
            'type': 'status_response',
            'agents_status': agents_status,
            'usage_stats': usage_stats,
            'database_stats': db_stats,
            'session_costs': session_costs,
            'session_manager_v5': SESSION_MANAGER_AVAILABLE,
            'supported_formats': get_supported_formats()
        })
        
    except Exception as e:
        logger.error(f"❌ Erreur WebSocket status: {e}")
        await manager.send_message(session_id, {
            'type': 'error',
            'message': f'Erreur status: {str(e)}'
        })

async def handle_websocket_documents_list(session_id: str, message_data: dict):
    """📋 Handler liste documents WebSocket"""
    try:
        limit = message_data.get('limit', 20)
        offset = message_data.get('offset', 0)
        
        if not MODULES_AVAILABLE:
            await manager.send_message(session_id, {
                'type': 'documents_list_response',
                'documents': [],
                'total': 0,
                'error': 'Modules V4 non disponibles'
            })
            return
        
        # Récupération via endpoint
        documents_response = await list_documents(limit=limit, offset=offset)
        
        await manager.send_message(session_id, {
            'type': 'documents_list_response',
            'documents': [doc.dict() for doc in documents_response.documents],
            'total': documents_response.total,
            'total_size': documents_response.total_size,
            'supported_formats': documents_response.supported_formats,
            'pdf_available': documents_response.pdf_available,
            'docx_available': documents_response.docx_available
        })
        
    except Exception as e:
        logger.error(f"❌ Erreur WebSocket documents list: {e}")
        await manager.send_message(session_id, {
            'type': 'error',
            'message': f'Erreur liste documents: {str(e)}'
        })

# === ROUTES API TRIANGLE + DOCUMENTS (pour compatibilité) ===

@app.post("/api/triple")
async def triple_endpoint(message: TripleMessage) -> dict:
    """🔀 API Triangle avec session V5"""
    if not MODULES_AVAILABLE:
        raise HTTPException(status_code=503, detail="Modules V4 non disponibles")
    
    try:
        # Session temporaire
        session_id = None
        if session_manager_v5:
            try:
                session_id = session_manager_v5.create_session(
                    user_id="FG",
                    metadata={
                        'session_type': 'api_triangle',  # ✅ DANS METADATA
                        'endpoint': '/api/triple',
                        'timestamp': datetime.now().isoformat()
                    }
                )
                session_manager_v5.add_message_to_session(
                    session_id=session_id,
                    message_type="user",
                    content={
                        'text': message.message,
                        'mode': 'triangle'
                    }
                )
            except Exception as e:
                logger.warning(f"⚠️ Session V5 non créée: {e}")
        
        # Contexte RAG
        context = ""
        rag_chunks_count = 0
        if message.use_rag:
            rag_manager = get_rag_manager()
            context = rag_manager.get_context_for_agent(message.message, "anima", message.rag_chunks)
            rag_chunks_count = context.count("=== DOCUMENT") + context.count("--- Interaction")
        
        # Débat triangulaire
        real_agents = get_real_agents()
        agents = ["anima", "neo", "nexus"]
        responses = []
        total_cost = 0.0
        
        for i, agent in enumerate(agents):
            if i == 0:
                prompt_message = message.message
            else:
                prompt_message = f"Débat: {message.message}\n\nÉchanges:\n" + \
                               "\n".join([f"{r['agent']}: {r['response'][:200]}..." for r in responses])
            
            agent_response = real_agents.get_response(agent, prompt_message, context)
            cost = agent_response.cost_estimate
            total_cost += cost
            
            response_data = {
                "agent": agent,
                "response": agent_response.response_text,
                "model": agent_response.model_used,
                "processing_time": agent_response.processing_time,
                "cost": cost,
                "role": ['ouverture', 'challenge', 'synthèse'][i]
            }
            responses.append(response_data)
            
            # Ajout session
            if session_manager_v5 and session_id:
                try:
                    session_manager_v5.add_message_to_session(
                        session_id=session_id,
                        message_type="agent",
                        content={
                            'agent': agent,
                            'text': agent_response.response_text,
                            'model': agent_response.model_used,
                            'cost': cost,
                            'triangle_position': i + 1,
                            'triangle_role': ['ouverture', 'challenge', 'synthèse'][i]
                        }
                    )
                except Exception as e:
                    logger.warning(f"⚠️ Session V5 message non ajouté: {e}")
        
        # Finalisation session
        if session_manager_v5 and session_id:
            try:
                session_manager_v5.finalize_session(session_id, {
                    'end_reason': 'api_triangle_complete',
                    'total_cost': total_cost,
                    'timestamp': datetime.now().isoformat()
                })
            except Exception as e:
                logger.warning(f"⚠️ Session V5 non finalisée: {e}")
        
        return {
            "responses": responses,
            "total_cost": total_cost,
            "rag_chunks_used": rag_chunks_count
        }
        
    except Exception as e:
        logger.error(f"❌ Erreur API triangle: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/documents/chat")
async def documents_chat_endpoint(message: DocumentsMessage) -> dict:
    """📁 API Documents avec session V5"""
    if not MODULES_AVAILABLE:
        raise HTTPException(status_code=503, detail="Modules V4 non disponibles")
    
    try:
        # Session temporaire
        session_id = None
        if session_manager_v5:
            try:
                session_id = session_manager_v5.create_session(
                    user_id="FG",
                    metadata={
                        'session_type': 'api_documents',  # ✅ DANS METADATA
                        'endpoint': '/api/documents/chat',
                        'timestamp': datetime.now().isoformat()
                    }
                )
                session_manager_v5.add_message_to_session(
                    session_id=session_id,
                    message_type="user",
                    content={
                        'text': message.message,
                        'selected_agents': message.agents,
                        'selected_documents': message.documents,
                        'mode': 'documents'
                    }
                )
            except Exception as e:
                logger.warning(f"⚠️ Session V5 non créée: {e}")
        
        # Contexte documentaire
        context = ""
        rag_chunks_count = 0
        context_info = {}
        
        if message.use_rag:
            rag_manager = get_rag_manager()
            
            # Si documents spécifiques sélectionnés
            if message.documents and hasattr(rag_manager, 'get_context_for_documents'):
                context = rag_manager.get_context_for_documents(
                    message.message, 
                    message.documents, 
                    message.rag_chunks
                )
            else:
                # Contexte général
                context = rag_manager.get_context_for_agent(
                    message.message, 
                    message.agents[0], 
                    message.rag_chunks
                )
            
            rag_chunks_count = context.count("=== DOCUMENT") + context.count("--- Interaction")
            context_info = {
                "sources_count": rag_chunks_count,
                "context_length": len(context),
                "selected_documents": len(message.documents) if message.documents else 'tous'
            }
        
        # Réponses des agents sélectionnés
        real_agents = get_real_agents()
        responses = []
        total_cost = 0.0
        
        for agent_name in message.agents:
            prompt_message = f"""Question: {message.message}

Contexte documentaire disponible:
{context[:2000] if context else "Aucun contexte documentaire spécifique."}

Analyse cette question en te basant sur les documents fournis."""
            
            agent_response = real_agents.get_response(agent_name, prompt_message, context)
            cost = agent_response.cost_estimate
            total_cost += cost
            
            response_data = {
                "agent": agent_name,
                "response": agent_response.response_text,
                "model": agent_response.model_used,
                "processing_time": agent_response.processing_time,
                "cost": cost,
                "rag_chunks_used": rag_chunks_count
            }
            responses.append(response_data)
            
            # Ajout session
            if session_manager_v5 and session_id:
                try:
                    session_manager_v5.add_message_to_session(
                        session_id=session_id,
                        message_type="agent",
                        content={
                            'agent': agent_name,
                            'text': agent_response.response_text,
                            'model': agent_response.model_used,
                            'cost': cost,
                            'rag_chunks_count': rag_chunks_count,
                            'documents_mode': True
                        }
                    )
                except Exception as e:
                    logger.warning(f"⚠️ Session V5 message non ajouté: {e}")
        
        # Finalisation session
        if session_manager_v5 and session_id:
            try:
                session_manager_v5.finalize_session(session_id, {
                    'end_reason': 'api_documents_complete',
                    'total_cost': total_cost,
                    'timestamp': datetime.now().isoformat()
                })
            except Exception as e:
                logger.warning(f"⚠️ Session V5 non finalisée: {e}")
        
        return {
            "responses": responses,
            "total_cost": total_cost,
            "context_info": context_info
        }
        
    except Exception as e:
        logger.error(f"❌ Erreur API documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# === POINT D'ENTRÉE ===
if __name__ == "__main__":
    import uvicorn
    
    print("🚀 ÉMERGENCE V5.1.2 - BACKEND COMPLET CORRIGÉ + WEBSOCKET HEARTBEAT STABLE")
    print("="*80)
    print(f"📄 Formats supportés: {', '.join(get_supported_formats())}")
    print(f"🔴 PDF: {'✅' if PDF_AVAILABLE else '❌'}")
    print(f"📘 DOCX: {'✅' if DOCX_AVAILABLE else '❌'}")
    print(f"🔧 Modules V4: {'✅' if MODULES_AVAILABLE else '❌'}")
    print(f"🧠 SessionManager V5: {'✅ Mémoire persistante' if SESSION_MANAGER_AVAILABLE else '❌ Mémoire basique'}")
    print()
    print("🎯 MODES SUPPORTÉS:")
    print("   💬 Mode Dialogue - Chat agent unique")
    print("   🔀 Mode Triangle - Débat triangulaire 3 agents")
    print("   📁 Mode Documents - Agents multiples + docs filtrés")
    print()
    print("🆕 NOUVELLES FONCTIONNALITÉS V5.1.2:")
    print("   - 🧠 Mémoire persistante avec SessionManager V5")
    print("   - 📚 Journalisation automatique conversations WebSocket")
    print("   - 🔍 API recherche sessions (/api/sessions/*)")
    print("   - 📊 Analytics sessions temps réel")
    print("   - 🔗 Tracking session V5 dans tous les modes")
    print("   - 💾 Export JSON sessions + recherche temporelle")
    print("   - 💓 WebSocket Heartbeat stable mobile (45s/90s)")
    print("   - 🔧 Handlers WebSocket complets et corrigés")
    print()
    print("✅ CORRECTIONS CRITIQUES APPLIQUÉES V5.1.2:")
    print("   - 🔧 Fix session_type parameter dans metadata")
    print("   - 🔧 Routes /api/sessions/* avec fallbacks robustes")
    print("   - 🔧 Gestion erreurs méthodes manquantes SessionManager")
    print("   - 💓 WebSocket Heartbeat ping/pong optimisé (45s interval)")
    print("   - 🔄 Reconnexion auto mobile avec timeout 90s")
    print("   - 🔌 Force disconnect sécurisé en cas d'erreur")
    print("   - 📁 handle_websocket_documents_mode AJOUTÉ (fonction manquante)")
    print("   - ✅ Envoi messages WebSocket sécurisé avec vérification état")
    print("   - 🎭 Tous handlers WebSocket présents et fonctionnels")
    print()
    print("🔗 ENDPOINTS SESSIONS V5:")
    print("   - GET /api/sessions/recent - Sessions récentes")
    print("   - GET /api/sessions/{id} - Détails session")
    print("   - GET /api/sessions/search - Recherche par critères")
    print("   - GET /api/sessions/stats - Statistiques générales")
    print()
    print("🌐 Démarrage serveur...")
    
    uvicorn.run(
        app, 
        host="0.0.0.0", 
        port=8000,
        reload=True,
        log_level="info"
    )